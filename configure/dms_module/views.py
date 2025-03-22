from django.shortcuts import render
from rest_framework import viewsets, filters
from rest_framework.response import Response
from .models import *
from rest_framework.pagination import PageNumberPagination
from .serializers import *
from rest_framework import permissions
from user_profile.function_call import *
from django.contrib.auth.models import Group
from django.core.mail import send_mail
from user_profile.email_utils import *
from lms_module.models import *
import requests
from django.db.models import Q
import ipdb
import logging
import time
import openpyxl
from django.core.files.base import ContentFile
from uuid import uuid4
from openpyxl.utils import get_column_letter
from io import BytesIO
from django.template.loader import get_template
from xhtml2pdf import pisa
from django.conf import settings

logger = logging.getLogger(__name__)

class CustomPagination(PageNumberPagination):
    page_size = 10  # Number of items per page
    page_size_query_param = 'page_size'  # Allow users to set page size
    max_page_size = 100  # Maximum page size


class DashboardCountViewSet(viewsets.ViewSet):
    permission_classes = [permissions.IsAuthenticated]

    def list(self, request, *args, **kwargs):
        try:
            document_count = Document.objects.count()
            workflow_count = WorkFlowModel.objects.count()
            document_type_count = DocumentType.objects.count()
            user_count = CustomUser.objects.count()
            department_count = Department.objects.count()

            return Response({
                "status": True,
                "message": "Dashboard counts fetched successfully",
                "data": {
                    "document_count": document_count,
                    "workflow_count": workflow_count,
                    "document_type_count": document_type_count,
                    "user_count" : user_count,
                    "department_count" : department_count
                }
            })
        except Exception as e:
            return Response({
                "status": False,
                "message": "Something went wrong",
                "error": str(e)
            })

class WorkFlowViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = WorkFlowSerializer
    queryset = WorkFlowModel.objects.all().order_by('-id')
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['workflow_name', 'workflow_description']
    ordering_fields = ['workflow_name', 'created_at']

    def create(self, request):
        try:
            workflow_name = request.data.get('workflow_name')
            workflow_description = request.data.get('workflow_description', '')
            is_active = request.data.get('is_active', True)

            if not workflow_name:
                return Response({'status': False, 'message': 'Workflow name is required'})

            workflow_obj = WorkFlowModel.objects.create(
                workflow_name=workflow_name,
                workflow_description=workflow_description,
                # is_active=is_active
            )
            serializer = WorkFlowSerializer(workflow_obj)
            return Response({'status': True, 'message': 'Workflow created successfully', 'data': serializer.data})
        except Exception as e:
            return Response({'status': False, 'message': 'Something went wrong', 'error': str(e)})

    # List all workflows
    def list(self, request):
        queryset = WorkFlowModel.objects.all().order_by('-id')
        
        try:
            if queryset.exists():
                serializer = WorkFlowSerializer(queryset, many=True)
                return Response({
                    "status": True,
                    "message": "Workflows fetched successfully",
                    'total': queryset.count(),
                    'data': serializer.data
                })
            else:
                return Response({
                    "status": True,
                    "message": "No Workflow found",
                    "total": 0,
                    "data": []
                })
        except Exception as e:
            return Response({"status": False, 'message': 'Something went wrong', 'error': str(e)})


class WorkFlowUpdateSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = WorkFlowSerializer
    lookup_field = 'workflow_id'

    def update(self, request, *args, **kwargs):
        try:
            workflow_id = self.kwargs.get("workflow_id")
            workflow_name = request.data.get('workflow_name')
            workflow_description = request.data.get('workflow_description')
    
            workflow_object = WorkFlowModel.objects.get(id=workflow_id)
            if workflow_name:
                workflow_object.workflow_name = workflow_name
            if workflow_description:
                workflow_object.workflow_description = workflow_description
            workflow_object.save()
    
            return Response({'status': True, 'message': 'Workflow updated successfully'})
        except WorkFlowModel.DoesNotExist:
            return Response({'status': False, 'message': 'Workflow not found'})
        except Exception as e:
            return Response({'status': False, 'message': 'Something went wrong', 'error': str(e)})
        
    def destroy(self, request, *args, **kwargs):
        try:
            workflow_id = request.data.get('workflow_id')   
            if not WorkFlowModel.objects.filter(id=workflow_id):
                return Response({"status":False, "message":"Workflow id not found"})
                     
            department_object = WorkFlowModel.objects.get(id=workflow_id)
            department_object.delete()
            return Response({"status":True, "message":"Workflow deleted succesfully"})
        except Exception as e:
                return Response({"status": False,'message': 'Something went wrong','error': str(e)})



class DocumentTypeCreateViewSet(viewsets.ModelViewSet):
    serializer_class = DocumentTypeSerializer
    queryset = DocumentType.objects.all()
    permission_classes = [permissions.IsAuthenticated]
    
    def create(self, request, *args, **kwargs):
        try:
            user = self.request.user
            document_name = request.data.get('document_name')
            if not document_name:
                return Response({"status": False, "message": "Document name is required", "data": []})
            document_type = DocumentType.objects.create(user = user, document_name=document_name)
            serializer = DocumentTypeSerializer(document_type)
            return Response({"status": True, "message": "Document type created successfully", "data": serializer.data})
        except Exception as e:
            return Response({"status": False, "message": str(e), "data": []})
        
    def list(self, request, *args, **kwargs):
        try:
            queryset = DocumentType.objects.all()
            serializer = DocumentTypeSerializer(queryset, many=True)
            return Response({"status": True, "message": "Document type list fetched successfully", "data": serializer.data})
        except Exception as e:
            return Response({"status": False, "message": str(e), "data": []})
                
class PrintRequestViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = PrintRequest.objects.all().order_by('-id')

    def create(self, request):
        try:
            user = self.request.user
            sop_document_id = request.data.get('sop_document_id')
            no_of_print = request.data.get('no_of_print')
            issue_type = request.data.get('issue_type')
            reason_for_print = request.data.get('reason_for_print')
            printer_id = request.data.get('printer_id')
            dynamic_status = DynamicStatus.objects.get(id=13)

            if not no_of_print:
                return Response({'status': False, 'message': 'NO of print is required'})
            if not issue_type:
                return Response({'status': False, 'message': 'Issue type is required'})
            if not reason_for_print:
                return Response({'status': False, 'message': 'Reason is required'})
            if not printer_id:
                return Response({'status': False, 'message': 'Printer id is required'})
            
            try:
                sop_document = Document.objects.get(id=sop_document_id)
            except Document.DoesNotExist:
                return Response({'status': False, 'message': 'Invalid document id'})
            
            try:
                printer = PrinterMachinesModel.objects.get(id=printer_id)
            except Document.DoesNotExist:
                return Response({'status': False, 'message': 'Invalid printer id'})

            printrequest_obj = PrintRequest.objects.create(
                user = user,
                no_of_print=no_of_print,
                issue_type=issue_type,
                reason_for_print=reason_for_print,
                sop_document_id = sop_document,
                printer = printer,
                print_request_status = dynamic_status
            )
            user_department = user.department
            qa_group = Group.objects.get(name='Doc Admin')
            qa_users_in_department = CustomUser.objects.filter(groups=qa_group, department=user_department)
            send_print_request_email(user, no_of_print, reason_for_print, sop_document, issue_type, qa_users_in_department)
            return Response({'status': True, 'message': 'Print requested successfully'})
        except Exception as e:
            return Response({'status': False, 'message': 'Something went wrong', 'error': str(e)})
        
    def list(self, request, *args, **kwargs):
        try:
            status_id = request.query_params.get('status_id', None)
            if not status_id:
                return Response({"status": False, "message": "status_id parameter is required", "data": []})
            user = self.request.user

            if user.groups.filter(name='QA').exists() or user.groups.filter(name='Doc Admin').exists():
                queryset = PrintRequest.objects.filter(print_request_status__id=13).order_by('-created_at')
            elif status_id == 'all':
                # If status_id is "all", fetch all data for the current user
                queryset = PrintRequest.objects.filter(user=user).order_by('-created_at')

            else:
                try:
                    dynamic_status = DynamicStatus.objects.get(id=status_id)
                except DynamicStatus.DoesNotExist:
                    return Response({"status": False, "message": "DynamicStatus with the given ID does not exist.", "data": []})

                queryset = PrintRequest.objects.filter(user=user, print_request_status=dynamic_status).order_by('-created_at')

            if not queryset:
                return Response({"status": False, "message": "No data available for the selected status.", "data": []})
            serializer = PrintRequestSerializer(queryset, many=True, context={'request': request})
            return Response({"status": True,"message": "Print requests fetched successfully.","total": queryset.count(),"data": serializer.data})

        except Exception as e:
            return Response({"status": False, "message": str(e), "data": []})

import os
import subprocess
from io import BytesIO
from django.conf import settings
from rest_framework import viewsets, permissions
from rest_framework.response import Response
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from dms_module.models import PrintRequest
from dms_module.serializers import PrintRequestSerializer

class PrintRequestDocxConvertPDFViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = PrintRequest.objects.all().order_by('-created_at')
    serializer_class = PrintRequestSerializer

    def update(self, request, *args, **kwargs):
        try:
            sop_document_id = self.kwargs.get('sop_document_id')
            document_status = request.data.get('document_status')
            approval_numbers = request.data.get('approval_numbers', [])
            sop_document_instance = PrintRequest.objects.filter(sop_document_id=sop_document_id).first()

            if not sop_document_instance:
                return Response({'status': False, 'message': 'Print request not found.'})

            # 🔹 Increase print count
            sop_document_instance.print_count += 1
            sop_document_instance.save()
            print_count = sop_document_instance.print_count - 1 

            # 🔹 Select the correct approval number
            if print_count < len(approval_numbers):
                approval_number = approval_numbers[print_count]
            else:
                approval_number = approval_numbers[-1] if approval_numbers else ""

            sop_document_file = sop_document_instance.sop_document_id
            sop_document = sop_document_file.generatefile
            print('Sop document:', sop_document)

            if not sop_document.endswith('.docx'):
                return Response({'status': False, 'message': 'Invalid document type. Only .docx files are supported.'})

            base_directory = os.path.join(settings.MEDIA_ROOT, 'generated_docs')
            docx_file_path = os.path.join(base_directory, sop_document)
            pdf_output_path = docx_file_path.replace('.docx', '.pdf')

            if not os.path.exists(docx_file_path):
                return Response({'status': False, 'message': f"Document file not found at {docx_file_path}."})

            try:
                # 🔹 Always Convert DOCX to PDF
                libreoffice_path = "libreoffice"
                if os.name == "nt":
                    libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"

                command = [libreoffice_path, "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(pdf_output_path), docx_file_path]
                subprocess.run(command, check=True)
            except subprocess.CalledProcessError as e:
                return Response({'status': False, 'message': 'Error during conversion.', 'error': str(e)})

            # 🔹 Apply Watermark Every Time
            watermarked_pdf_path = pdf_output_path
            add_watermark(pdf_output_path, watermarked_pdf_path, document_status, approval_number)
            restrict_pdf_print(watermarked_pdf_path)
            pdf_relative_path = os.path.relpath(watermarked_pdf_path, settings.MEDIA_ROOT)
            pdf_url = f"{settings.MEDIA_URL}{pdf_relative_path}"

            return Response({'status': True, 'message': 'Document successfully converted with watermark.', 'pdf_link': request.build_absolute_uri(pdf_url)})

        except Exception as e:
            return Response({'status': False, 'message': 'An error occurred while processing the document.', 'error': str(e)})

from reportlab.lib.colors import black, lightgrey
def add_watermark(input_pdf_path, output_pdf_path, watermark_text="CONFIDENTIAL", approval_number=""):
    """ Adds a diagonal watermark and places the approval number at the bottom-right of each page with improvements """

    existing_pdf = PdfReader(open(input_pdf_path, "rb"))
    output = PdfWriter()

    for i in range(len(existing_pdf.pages)):
        page = existing_pdf.pages[i]

        # Create a new PDF for watermark & approval number
        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        width, height = letter  # Get PDF page size dynamically

        # 🔹 Add Semi-Transparent Watermark (Diagonal)
        can.saveState()
        can.setFont("Helvetica-Bold", 50)
        can.setFillColorRGB(0.8, 0.1, 0.1, alpha=0.3)  # Red, semi-transparent
        can.translate(width / 3, height / 2)
        can.rotate(45)
        can.drawString(0, 0, watermark_text)
        can.restoreState()  # Restore to avoid affecting approval number

        # 🔹 Add Black Approval Number at **Bottom-Right** of Each Page
        if approval_number:
            can.saveState()
            can.setFont("Courier-Bold", 14)  # Use "Courier-Bold" for better clarity
            can.setFillColor(black)  # Black color
            box_width, box_height = 140, 25  # Size of background box
            padding = 5

            # Draw background box for better visibility
            can.setFillColor(lightgrey)  # Light grey background
            can.rect(width - box_width - 10, 10, box_width, box_height, fill=1, stroke=0)

            # Add approval number text
            can.setFillColor(black)  # Reset text color to black
            can.drawString(width - box_width + padding - 10, 20, approval_number)

            can.restoreState()

        can.save()
        packet.seek(0)
        watermark_pdf = PdfReader(packet)

        # Merge the watermark layer with the existing page
        page.merge_page(watermark_pdf.pages[0])
        output.add_page(page)

    # Save the final PDF
    with open(output_pdf_path, "wb") as outputStream:
        output.write(outputStream)


class PrintRequestExcelGenerateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = PrintRequest.objects.all().order_by('-created_at')
    serializer_class = PrintRequestSerializer

    def list(self, request, *args, **kwargs):
        try:
            status_id = request.query_params.get('status_id', None)
            if not status_id:
                return Response({"status": False, "message": "status_id parameter is required", "data": []})
            user = self.request.user

            if user.groups.filter(name='QA').exists() or user.groups.filter(name='Doc Admin').exists():
                queryset = PrintRequest.objects.filter(print_request_status__id=12).order_by('-created_at')

            elif status_id == 'all':
                # If status_id is "all", fetch all data for the current user
                queryset = PrintRequest.objects.filter(user=user).order_by('-created_at')

            else:
                try:
                    dynamic_status = DynamicStatus.objects.get(id=status_id)
                except DynamicStatus.DoesNotExist:
                    return Response({"status": False, "message": "DynamicStatus with the given ID does not exist.", "data": []})

                queryset = PrintRequest.objects.filter(user=user, print_request_status=dynamic_status).order_by('-created_at')

            if not queryset:
                return Response({"status": False, "message": "No data available for the selected status.", "data": []})

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Print Requests"

            headers = [
                'Data Number', 'User','SOP Document Number','SOP Document Title', 'No of Prints', 'Retrieval Numbers', 'Issue Type', 'Reason for Print',
                'Print Request Status', 'Created At', 'Printer', 'Master Copy Users', 'Other Users', 'Reminder Sent'
            ]

            for col_num, header in enumerate(headers, 1):
                col_letter = get_column_letter(col_num)
                ws[f'{col_letter}1'] = header

            for row_num, (index, print_request) in enumerate(enumerate(queryset, start=1), 2):
                ws[f'A{row_num}'] = index
                ws[f'B{row_num}'] = print_request.user.username  # User's username
                ws[f'C{row_num}'] = print_request.sop_document_id.document_number if print_request.sop_document_id else ""
                ws[f'D{row_num}'] = print_request.sop_document_id.document_title if print_request.sop_document_id else ""
                ws[f'E{row_num}'] = print_request.no_of_print
                ws[f'F{row_num}'] = ", ".join([str(num) for num in print_request.approvals.all().values_list('retrival_numbers__retrival_number', flat=True) if num])
                ws[f'G{row_num}'] = print_request.issue_type
                ws[f'H{row_num}'] = print_request.reason_for_print
                status = print_request.print_request_status.status if print_request.print_request_status else ""
                ws[f'I{row_num}'] = status.capitalize() if status else ""
                ws[f'J{row_num}'] = print_request.created_at.strftime('%d-%m-%Y')
                ws[f'K{row_num}'] = print_request.printer.printer_name if print_request.printer else ""
                ws[f'L{row_num}'] = ", ".join([user.username for user in print_request.master_copy_user.all()])
                ws[f'M{row_num}'] = ", ".join([user.username for user in print_request.other_user.all()])
                ws[f'N{row_num}'] = "Yes" if print_request.reminder_sent else "No"

            for col_num in range(1, len(headers) + 1):
                col_letter = get_column_letter(col_num)
                max_length = 0
                for row in ws.iter_rows(min_col=col_num, max_col=col_num):
                    for cell in row:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(cell.value)
                        except:
                            pass
                adjusted_width = (max_length + 2)
                ws.column_dimensions[col_letter].width = adjusted_width
            timestamp = time.strftime("%d_%m_%Y_%H_%M_%S")
            filename = f"print_request_report_{timestamp}.xlsx"

            file_path = os.path.join(settings.MEDIA_ROOT, 'print_request_excel_sheet', filename)

            folder_path = os.path.dirname(file_path)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)

            file_stream = BytesIO()
            wb.save(file_stream)
            file_stream.seek(0)

            with open(file_path, 'wb') as f:
                f.write(file_stream.read())
            base_url = request.build_absolute_uri('/')
            file_url = base_url + 'media/print_request_excel_sheet/' + filename

            return Response({"status": True,"message": "Excel report generated successfully.","data": file_url})

        except Exception as e:
            return Response({"status": False, "message": str(e), "data": []})

# class PrintApprovalViewSet(viewsets.ModelViewSet):
#     permission_classes = [permissions.IsAuthenticated]
#     queryset = PrintRequestApproval.objects.all().order_by('-id')

#     def create(self, request, *args, **kwargs):
#         try:
#             user = self.request.user
#             print_request_id = request.data.get('print_request_id')
#             no_of_request_by_admin = request.data.get('no_of_request_by_admin')
#             status = request.data.get('status')

#             # Validate that the print_request_id is provided
#             if not print_request_id:
#                 return Response({'status': False, 'message': 'Print request ID is required'})
            
#             # Fetch the associated PrintRequest object
#             try:
#                 print_request = PrintRequest.objects.get(id=print_request_id)
#             except PrintRequest.DoesNotExist:
#                 return Response({'status': False, 'message': 'Invalid Print Request ID'})
            
#             if not no_of_request_by_admin:
#                 return Response({'status': False, 'message': 'No of request by admin is required when approving'})
            
#             # Ensure no_of_request_by_admin does not exceed no_of_print from PrintRequest
#             if int(no_of_request_by_admin) > print_request.no_of_print:
#                 return Response({'status': False, 'message': f'No of request by admin cannot exceed {print_request.no_of_print}'})
            
#             # Fetch the DynamicStatus object for the provided status
#             try:
#                 dynamic_status = DynamicStatus.objects.get(id=status)
#             except DynamicStatus.DoesNotExist:
#                 return Response({'status': False, 'message': 'Invalid status value provided'})
            
#             # Generate the unique approval number
#             base_format = "BPL-Dms-"

#             # Track the last approval for the current print_request
#             last_approval = PrintRequestApproval.objects.filter(print_request=print_request).order_by('-id').first()

#             # Initialize the copy number for the first approval
#             if last_approval:
#                 # Extract the last copy number and increment it
#                 last_copy_number = int(last_approval.approval_number.split('-')[-2])
#                 new_copy_number = last_copy_number + 1
#             else:
#                 # If no approval exists yet, start from 01
#                 new_copy_number = 1

#             # Ensure that the generated approval number is unique
#             approval_number = f"{base_format}{str(new_copy_number).zfill(2)}-{no_of_request_by_admin}"

#             # Check if the generated approval number already exists, and keep incrementing until it's unique
#             while PrintRequestApproval.objects.filter(approval_number=approval_number).exists():
#                 new_copy_number += 1
#                 approval_number = f"{base_format}{str(new_copy_number).zfill(2)}-{no_of_request_by_admin}"

#             # Create PrintRequestApproval object
#             print_request_approval = PrintRequestApproval.objects.create(
#                 user=user,
#                 print_request=print_request,
#                 no_of_request_by_admin=no_of_request_by_admin,
#                 status=dynamic_status,
#                 approval_number=approval_number
#             )
            
#             return Response({'status': True, 'message': f'Print request successfully'})
        
#         except Exception as e:
#             return Response({'status': False, 'message': 'Something went wrong', 'error': str(e)})

class PrintApprovalViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = PrintRequestApproval.objects.all().order_by('-id')

    def create(self, request, *args, **kwargs):
        try:
            user = self.request.user
            print_request_id = request.data.get('print_request_id')
            no_of_request_by_admin = request.data.get('no_of_request_by_admin')
            status = request.data.get('status')

            # Validate that the print_request_id is provided
            if not print_request_id:
                return Response({'status': False, 'message': 'Print request ID is required'})
            
            # Fetch the associated PrintRequest object
            try:
                print_request = PrintRequest.objects.get(id=print_request_id)
            except PrintRequest.DoesNotExist:
                return Response({'status': False, 'message': 'Invalid Print Request ID'})
            
            if not no_of_request_by_admin:
                return Response({'status': False, 'message': 'No of request by admin is required when approving'})
            
            # Ensure no_of_request_by_admin does not exceed no_of_print from PrintRequest
            if int(no_of_request_by_admin) > print_request.no_of_print:
                return Response({'status': False, 'message': f'No of request by admin cannot exceed {print_request.no_of_print}'})
            
            # Fetch the DynamicStatus object for the provided status
            try:
                dynamic_status = DynamicStatus.objects.get(id=status)
            except DynamicStatus.DoesNotExist:
                return Response({'status': False, 'message': 'Invalid status value provided'})
            
            # Generate unique approval numbers
            base_format = "BPL-Dms-"
            approval_numbers = []
            for i in range(int(no_of_request_by_admin)):
                new_copy_number = 1  # Start from 1
                approval_number = f"{base_format}{str(new_copy_number).zfill(2)}-{i+1}"

                # Ensure uniqueness in the database
                while ApprovalNumber.objects.filter(number=approval_number).exists():
                    new_copy_number += 1
                    approval_number = f"{base_format}{str(new_copy_number).zfill(2)}-{i+1}"

                # Create or get the unique approval number
                approval_number_obj, _ = ApprovalNumber.objects.get_or_create(number=approval_number)
                approval_numbers.append(approval_number_obj)

            # Create PrintRequestApproval object
            print_request_approval = PrintRequestApproval.objects.create(
                user=user,
                print_request=print_request,
                no_of_request_by_admin=no_of_request_by_admin,
                status=dynamic_status,
            )
            # Add approval numbers to the ManyToManyField
            print_request_approval.approval_numbers.add(*approval_numbers)
            print_request_approval.save()
            print_request.print_request_status = dynamic_status
            print_request.save()
            #send email to the user who requested the print in PrintRequest
            user = print_request.user
            send_print_request_approval_email(user, print_request, no_of_request_by_admin, dynamic_status)
            return Response({
                'status': True,
                'message': 'Print request successfully approved',
            })
        
        except Exception as e:
            return Response({'status': False, 'message': 'Something went wrong', 'error': str(e)})

        

class PrintRequestUpdateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]

    def update(self, request, *args, **kwargs):
        try:
            print_request_id = self.kwargs.get('print_request_id')
            status = request.data.get('status')

            if not print_request_id:
                return Response({'status': False, 'message': 'Print request ID is required'})

            try:
                print_request = PrintRequest.objects.get(id=print_request_id)
            except PrintRequest.DoesNotExist:
                return Response({'status': False, 'message': 'Invalid Print Request ID'})

            try:
                print_request_approval = PrintRequestApproval.objects.get(print_request=print_request, status='approved')
            except PrintRequestApproval.DoesNotExist:
                return Response({'status': False, 'message': 'No approved PrintRequestApproval found for this PrintRequest'})

            if print_request.print_request_status != 'print_is_pending':
                return Response({'status': False, 'message': 'This PrintRequest is not in a pending state'})

            print_request.print_request_status = status
            print_request.save()

            return Response({'status': True, 'message': 'Data Printing started successfully'})
        
        except Exception as e:
            return Response({'status': False, 'message': 'Something went wrong', 'error': str(e)})


class DocumentCreateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentdataSerializer
    queryset = Document.objects.all()
    
    def create(self, request, *args, **kwargs):
        try:
            user = self.request.user
            parent_document = request.data.get('parent_document', None)
            equipment_id = request.data.get('equipment_id','')
            product_code = request.data.get('product_code','')
            document_title = request.data.get('document_title')
            document_type_id = request.data.get('document_type')
            document_description = request.data.get('document_description', '')
            revision_month = request.data.get('revision_month', '')
            document_operation = request.data.get('document_operation', '')
            select_template = request.data.get('select_template')
            workflow = request.data.get('workflow')
            document_current_status_id = request.data.get('document_current_status_id')
            training_required = request.data.get('training_required')
            if training_required != '':
                if isinstance(training_required, str):
                    training_required = training_required.lower() in ['true',1]
            else:
                training_required = False
            # visible_to_users = request.data.get('visible_to_users', [])
            # approver = request.data.get('approver')
            # doc_admin = request.data.get('doc_admin')


             # Ensure visible_to_users is parsed into a proper list
            # if isinstance(visible_to_users, str):
            #     import json
            #     try:
            #         visible_to_users = json.loads(visible_to_users)
            #     except json.JSONDecodeError:
            #         return Response({
            #             "status": False,
            #             "message": "Invalid format for visible_to_users. Provide a valid list of IDs.",
            #             "data": []
            #         })

            # Validation for required fields
            if not document_title:
                return Response({"status": False, "message": "Document title is required", "data": []})
            if not document_type_id:
                return Response({"status": False, "message": "Document type is required", "data": []})
            if not workflow:
                return Response({"status": False, "message": "Workflow is required", "data": []})
            if not select_template:
                    return Response({"status": False, "message": "Template selection is required", "data": []})
            if not revision_month:
                return Response({"status": False, "message": "Revision month is required", "data": []})
            # if not approver:
            #     return Response({"status": False, "message": "Please select Approver", "data": []})
            # if not doc_admin:
            #     return Response({"status": False, "message": "Please select Doc Admin", "data": []})

            try:
                default_status = DynamicStatus.objects.get(id=document_current_status_id)  # Assuming status with ID 1 is the default
            except DynamicStatus.DoesNotExist:
                return Response({"status": False, "message": "Default status not found in the system", "data": []})

            # Fetch DocumentType instance
            try:
                document_type = DocumentType.objects.get(id=document_type_id)
            except DocumentType.DoesNotExist:
                return Response({"status": False, "message": "Document type not found", "data": []})
            
            # try:
            #     approver_user = CustomUser.objects.get(id=approver)
            # except DocumentType.DoesNotExist:
            #     return Response({"status": False, "message": "Approver user not found", "data": []})
            
            # try:
            #     docadmin_user = CustomUser.objects.get(id=doc_admin)
            # except DocumentType.DoesNotExist:
            #     return Response({"status": False, "message": "Doc Admin user not found", "data": []})

            # Handle Parent Document if provided
            # parent_document_instance = []
            # if parent_document:
            #     try:
            #         parent_document_instance = Document.objects.get(id__in=parent_document)
            #         parent_document_instance.is_parent = True
            #         parent_document_instance.save()
            #     except Document.DoesNotExist:
            #         return Response({"status": False, "message": "Parent document not found", "data": []})
            parent_document_instance = None
            if parent_document:
                    try:
                        parent_document_instance = Document.objects.filter(id=parent_document).first()
                        parent_document_instance.is_parent = True
                        parent_document_instance.save()
                    except Document.DoesNotExist:
                        return Response({"status": False, "message": "Parent document not found", "data": []})
                
            department_name = user.department.department_name if user.department else 'UnknownDepartment'
            document_number = generate_document_number(user, document_type, parent_document_instance)

            document = Document.objects.create(
                user=user,
                parent_document = parent_document_instance,
                document_title=document_title,
                document_number=document_number,
                document_type_id=document_type.id,
                document_description=document_description,
                product_code=product_code,
                equipment_id=equipment_id,
                # revision_date=revision_date,
                revision_month = revision_month,
                document_operation=document_operation,
                select_template_id=select_template,
                workflow_id=workflow,
                document_current_status=default_status,
                version="0.0",
                # approver=approver_user,
                # doc_admin=docadmin_user,
                training_required=training_required,

            )
            # if visible_to_users:
            #     document.visible_to_users.set(visible_to_users)

            document.save()
            DocumentVersion.objects.create(
                document=document,
                version_number=document.version,
                updated_by=user,
            )

            # reviewer_group = Group.objects.get(name='Reviewer')
            # reviewers = CustomUser.objects.filter(groups=reviewer_group)
            # department_users = CustomUser.objects.filter(department=user.department)
            # users_to_notify = reviewers.union(department_users).distinct()
            # send_document_create_email(user, document_title, users_to_notify)

            # If the operation is 'upload_file', handle the Word file upload
            # if document_operation == 'upload_file':
            #     word_file = request.FILES['word_file']
            #     UploadedDocument.objects.create(
            #         document=document,
            #         word_file=word_file
            #     )
            
            # Serialize the created document
            # serializer = DocumentdataSerializer(document)
            return Response({"status": True, "message": "Document created successfully"})
        
        except Exception as e:
            return Response({"status": False, "message": str(e), "data": []})

class DocumentUpdateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentSerializer
    queryset = Document.objects.all()
    lookup_field = 'document_id'

    def update(self, request, *args, **kwargs):
        try:
            user = self.request.user
            document_id = self.kwargs.get('document_id')
            
            try:
                document = Document.objects.get(id=document_id)
            except Document.DoesNotExist:
                return Response({'status': False, 'message': 'Document not found'})

            # Extract fields from the request data
            parent_document = request.data.get('parent_document',None)
            document_title = request.data.get('document_title')
            document_type_id = request.data.get('document_type')
            document_description = request.data.get('document_description')
            revision_month = request.data.get('revision_month', '')
            document_operation = request.data.get('document_operation')
            workflow_id = request.data.get('workflow')
            training_required = request.data.get('training_required')
            select_template = request.FILES.get('select_template')
            equipment_id = request.data.get('equipment_id','')
            product_code = request.data.get('product_code','')
            # visible_to_users = request.data.get('visible_to_users', [])
            # approver = request.data.get('approver')
            # doc_admin = request.data.get('doc_admin')
            
            # Validate and parse visible_to_users
            # if isinstance(visible_to_users, str):
            #     import json
            #     try:
            #         visible_to_users = json.loads(visible_to_users)
            #     except json.JSONDecodeError:
            #         return Response({
            #             "status": False,
            #             "message": "Invalid format for visible_to_users. Provide a valid list of user IDs.",
            #             "data": []
            #         })

            if not document_title:
                return Response({"status": False, "message": "Document title is required", "data": []})
            if not document_type_id:
                return Response({"status": False, "message": "Document type is required", "data": []})
            if not workflow_id:
                return Response({"status": False, "message": "Workflow is required", "data": []})
            if not revision_month:
                return Response({"status": False, "message": "Revision month is required", "data": []})
            # if not approver:
            #     return Response({"status": False, "message": "Please select Approver", "data": []})
            # if not doc_admin:
            #     return Response({"status": False, "message": "Please select Doc Admin", "data": []})

            try:
                document_type = DocumentType.objects.get(id=document_type_id)
            except DocumentType.DoesNotExist:
                return Response({'status': False, 'message': 'Document type not found'})
            
            try:
                workflow = WorkFlowModel.objects.get(id=workflow_id)
            except WorkFlowModel.DoesNotExist:
                return Response({'status': False, 'message': 'Workflow not found'}, status=400)

            # try:
            #     approver_user = CustomUser.objects.get(id=approver)
            # except DocumentType.DoesNotExist:
            #     return Response({"status": False, "message": "Approver user not found", "data": []})
            
            # try:
            #     docadmin_user = CustomUser.objects.get(id=doc_admin)
            # except DocumentType.DoesNotExist:
            #     return Response({"status": False, "message": "Doc Admin user not found", "data": []})
      
            parent_document_instance = None
            if parent_document:
                try:
                    parent_document_instance = Document.objects.get(id=parent_document)
                except Document.DoesNotExist:
                    return Response({"status": False, "message": "Parent document not found", "data": []})

            if select_template:
                try:
                    # get_select_template = document.select_template.template_doc
                    document.select_template.template_doc = select_template
                    document.select_template.save()
                except Exception as e:
                    return Response({"status": False, "message": str(e), "data": []})
            # Update the document fields
            if document_title != '':
                document.document_title = document_title
            if parent_document != '':
                document.parent_document = parent_document_instance
            if document_type != '':
                document.document_type = document_type
            if document_description != '':
                document.document_description = document_description
            if revision_month != '':
                document.revision_month = revision_month
            if document_operation != '':
                document.document_operation = document_operation
            if workflow != '':
                document.workflow = workflow
            if equipment_id != '':
                document.equipment_id = equipment_id
            if product_code != '':
                document.product_code = product_code        
            if training_required != '':
                if isinstance(training_required, str):
                    training_required = training_required.lower() in ['true',1]
                document.training_required = training_required
            # if approver_user != '':
            #     document.approver = approver_user
            # if docadmin_user != '':
            #     document.doc_admin = docadmin_user 

            # # Update visible_to_users if provided
            # if visible_to_users:
            #     document.visible_to_users.set(visible_to_users)    
            
            # Save the updated document
            document.save()
            UpdateDocumentByUser.objects.create(user=user, document=document)

            return Response({"status": True, "message": "Document updated successfully"})
        
        except Exception as e:
            return Response({"status": False, "message": 'Something went wrong', 'error': str(e)})

# class DocumentViewSet(viewsets.ModelViewSet):
#     permission_classes = [permissions.IsAuthenticated]
#     serializer_class = DocumentviewSerializer
#     queryset = Document.objects.all().order_by('-id')
#     filter_backends = [filters.SearchFilter, filters.OrderingFilter]
#     search_fields = ['document_title', 'document_number', 'document_description', 'document_type__name']
#     ordering_fields = ['document_title', 'created_at'] 

#     def list(self, request):
#         try:
#             queryset = self.filter_queryset(self.get_queryset())

#             if queryset.exists():
#                 serializer = DocumentviewSerializer(queryset, many=True)
#                 return Response({
#                     "status": True,
#                     "message": "Documents fetched successfully",
#                     'total': queryset.count(),
#                     'data': serializer.data
#                 })
#             else:
#                 return Response({
#                     "status": True,
#                     "message": "No Documents found",
#                     "total": 0,
#                     "data": []
#                 })
#         except Exception as e:
#             return Response({"status": False, 'message': 'Something went wrong', 'error': str(e)})


# class DocumentViewSet(viewsets.ModelViewSet):
#     permission_classes = [permissions.IsAuthenticated]
#     serializer_class = DocumentviewSerializer
#     queryset = Document.objects.all().order_by('-id')
#     filter_backends = [filters.SearchFilter, filters.OrderingFilter]
#     search_fields = ['document_title', 'document_number', 'document_description', 'document_type__name']
#     ordering_fields = ['document_title', 'created_at'] 

#     def list(self, request):
#         try:
#             user = self.request.user 
#             user_group_ids = user.groups.values_list('id', flat=True)
#             user_department = user.department if user.department else None
#             department_id = request.query_params.get('department_id', None)
#             document_current_status = request.query_params.get('document_current_status', None)
#             start_date = request.query_params.get('start_date', None)
#             end_date = request.query_params.get('end_date', None)

#             start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
#             if error:
#                 return Response({"status": False, "message": error, "data": [],"user_group_ids": list(user_group_ids)})
            
#             if start_date_obj:
#                 start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
#             if end_date_obj:
#                 end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))
                
#             if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
#                 if department_id:
#                     queryset = Document.objects.filter(user__department_id=department_id).order_by('-id')
#                 else:
#                     queryset = Document.objects.all().order_by('-id')

#             elif user.groups.filter(name='Reviewer').exists():
#                 queryset = Document.objects.filter(visible_to_users=user).order_by('-id')
#             elif user_department:
#                 queryset = Document.objects.filter(user__department=user_department).order_by('-id')
#             else:
#                 queryset = Document.objects.none()

#             if start_date_obj and end_date_obj:
#                 queryset = queryset.filter(created_at__range=[start_date_obj, end_date_obj])

#             if document_current_status:
#                 # Filter by document_current_status if provided
#                 queryset = queryset.filter(document_current_status=document_current_status)
            
#             queryset = self.filter_queryset(queryset)

#             if queryset.exists():
#                 serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
#                 return Response({"status": True,"message": "Documents fetched successfully",'user_group_ids': list(user_group_ids) ,'data': serializer.data,})
#             else:
#                 return Response({"status": True,"message": "No Documents found",'user_group_ids': list(user_group_ids),"data": []})
        
#         except Exception as e:
#             return Response({"status": False,'message': 'Something went wrong','error': str(e)})



# class DocumentViewSet(viewsets.ModelViewSet):
#     permission_classes = [permissions.IsAuthenticated]
#     serializer_class = DocumentviewSerializer
#     queryset = Document.objects.all().order_by('-id')
#     filter_backends = [filters.SearchFilter, filters.OrderingFilter]
#     search_fields = ['document_title', 'document_number', 'document_description', 'document_type__name']
#     ordering_fields = ['document_title', 'created_at'] 

#     def list(self, request):
#         try:
#             user = self.request.user 
#             user_group_ids = user.groups.values_list('id', flat=True)
#             user_department = user.department if user.department else None
#             department_id = request.query_params.get('department_id', None)
#             document_current_status = request.query_params.get('document_current_status', None)

#             if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
#                 if department_id:
#                     queryset = Document.objects.filter(user__department_id=department_id).order_by('-id')
#                 else:
#                     queryset = Document.objects.all().order_by('-id')

#             elif user.groups.filter(name='Reviewer').exists():
#                 queryset = Document.objects.filter(visible_to_users=user).order_by('-id')
#             elif user_department:
#                 queryset = Document.objects.filter(user__department=user_department).order_by('-id')
#             else:
#                 queryset = Document.objects.none()

#             document_current_status = request.query_params.get('document_current_status', None)

#             if document_current_status:
#                 # Filter by document_current_status if provided
#                 queryset = queryset.filter(document_current_status=document_current_status)
            
#             queryset = self.filter_queryset(queryset)

#             if queryset.exists():
#                 serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
#                 return Response({"status": True,"message": "Documents fetched successfully",'user_group_ids': list(user_group_ids) ,'data': serializer.data,})
#             else:
#                 return Response({"status": True,"message": "No Documents found",'user_group_ids': list(user_group_ids),"data": []})
        
#         except Exception as e:
#             return Response({"status": False,'message': 'Something went wrong','error': str(e)})

class DocumentViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-created_at')
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['document_title', 'document_number', 'document_description', 'document_type__name']
    ordering_fields = ['document_title', 'created_at']

    def list(self, request):
        try:
            user = self.request.user
            user_group_ids = list(user.groups.values_list('id', flat=True))
            department_id = self.request.query_params.get('department_id', None)
            document_current_status = request.query_params.get('document_current_status', None)
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)

            start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
            if error:
                return Response({"status": False, "message": error, "data": [],"user_group_ids": list(user_group_ids)})
            
            if start_date_obj:
                start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
            if end_date_obj:
                end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))

             # Check if the user is in the "Admin" group
            if user.groups.filter(name="Admin").exists():
                # Admins can view all documents
                queryset = Document.objects.all().exclude(document_current_status=15).order_by('-id')
            elif user.groups.filter(name="Author").exists():
                # Authors can only see documents created by users in the same department
                if department_id:
                    queryset = Document.objects.filter(
                        user__department_id=department_id,
                        user__groups__name="Author"
                    ).exclude(document_current_status=15).order_by('-id') | Document.objects.filter(document_current_status=10) | Document.objects.filter(author=user).order_by('-id')
                else:
                    queryset = Document.objects.filter(
                        user__department_id=user.department_id,
                        user__groups__name="Author"
                    ).exclude(document_current_status=15).order_by('-id') | Document.objects.filter(document_current_status=10) | Document.objects.filter(author=user).order_by('-id')
            else:
                # Other roles: View documents assigned to them
                queryset = Document.objects.filter(
                    Q(approver=user) |
                    Q(doc_admin=user) |
                    Q(visible_to_users=user)
                ).exclude(document_current_status=15).distinct()

            # Filter by document_current_status if provided
            if document_current_status:
                queryset = queryset.filter(document_current_status=document_current_status)

            if start_date_obj and end_date_obj:
                queryset = queryset.filter(created_at__range=[start_date_obj, end_date_obj])

            queryset = self.filter_queryset(queryset)

            if queryset.exists():
                serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
                return Response({
                    "status": True,
                    "message": "Documents fetched successfully",
                    'user_group_ids': user_group_ids,
                    'data': serializer.data,
                })
            else:
                return Response({
                    "status": True,
                    "message": "No Documents found",
                    'user_group_ids': user_group_ids,
                    "data": []
                })

        except Exception as e:
            return Response({
                "status": False,
                'message': 'Something went wrong',
                'error': str(e)
            })








class DocumentExcelGenerateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['document_title', 'document_number', 'document_description', 'document_type__name']
    ordering_fields = ['document_title', 'created_at'] 

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user 
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            document_current_status = request.query_params.get('document_current_status', None)

            # Admin or Doc Admin: view all or by department
            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                if department_id:
                    queryset = Document.objects.filter(user__department_id=department_id).order_by('-id')
                else:
                    queryset = Document.objects.all().order_by('-id')

            # Reviewer: can only view documents visible to them
            elif user.groups.filter(name='Reviewer').exists():
                queryset = Document.objects.filter(visible_to_users=user).order_by('-id')

            # Department-based filtering for other users
            elif user_department:
                queryset = Document.objects.filter(user__department=user_department).order_by('-id')
            else:
                queryset = Document.objects.none()

            # Further filtering by document_current_status if provided
            if document_current_status:
                queryset = queryset.filter(document_current_status=document_current_status)

            # Apply any filters set in the queryset
            queryset = self.filter_queryset(queryset)

            # Generate the Excel file automatically
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Documents Report Excel Sheet"

            # Define the headers for the Excel file
            headers = [
                'Document Title', 'Document Number', 'Document Type', 'Parent Document No.',
                'Revision Date', 'Status','Assigned User FirstName','Assign User LastName','Created At', 'Effective Date', 'Version',
                'Author Name', 'Approver Name','Doc Admin Name'
            ]
            # Add headers to the Excel sheet
            for col_num, header in enumerate(headers, 1):
                col_letter = get_column_letter(col_num)
                ws[f'{col_letter}1'] = header

            # Add data rows from the queryset
            for row_num, document in enumerate(queryset, 2):
                ws[f'A{row_num}'] = document.document_title
                ws[f'B{row_num}'] = document.document_number
                ws[f'C{row_num}'] = document.document_type.document_name if document.document_type else "-"
                ws[f'D{row_num}'] = document.parent_document.id if document.parent_document else "-"
                ws[f'E{row_num}'] = document.revision_date.strftime('%d-%m-%Y') if document.revision_date else "-"
                ws[f'F{row_num}'] = document.document_current_status.status if document.document_current_status else "-"
                ws[f'G{row_num}'] = document.assigned_to.first_name if document.assigned_to else "-"
                ws[f'H{row_num}'] = document.assigned_to.last_name if document.assigned_to else "-"
                ws[f'I{row_num}'] = document.created_at.strftime('%d-%m-%Y')
                ws[f'J{row_num}'] = document.effective_date.strftime('%d-%m-%Y') if document.effective_date else "-"
                ws[f'K{row_num}'] = document.version
                ws[f'L{row_num}'] = document.author.username if document.author else "-"
                ws[f'M{row_num}'] = document.approver.username if document.approver else "-"
                ws[f'N{row_num}'] = document.doc_admin.username if document.doc_admin else "-"

                user_groups = document.user.groups.all()
                group_names = ', '.join([group.name for group in user_groups]) if user_groups else '-'
                # ws[f'L{row_num}'] = group_names

            # Adjust column widths
            for col_num in range(1, len(headers) + 1):
                col_letter = get_column_letter(col_num)
                max_length = 0
                for row in ws.iter_rows(min_col=col_num, max_col=col_num):
                    for cell in row:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(cell.value)
                        except:
                            pass
                adjusted_width = (max_length + 2)
                ws.column_dimensions[col_letter].width = adjusted_width

            # Generate a timestamp for the file name
            timestamp = time.strftime("%d_%m_%Y_%H_%M_%S")
            filename = f"document_report_{timestamp}.xlsx"

            # File path and saving the file
            file_path = os.path.join(settings.MEDIA_ROOT, 'document_excel_sheet', filename)

            # Create the folder if it doesn't exist
            folder_path = os.path.dirname(file_path)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)

            # Save the Excel file to the path
            file_stream = BytesIO()
            wb.save(file_stream)
            file_stream.seek(0)

            with open(file_path, 'wb') as f:
                f.write(file_stream.read())

            # Build the base URL and the file URL for downloading
            base_url = request.build_absolute_uri('/')
            file_url = base_url + 'media/document_excel_sheet/' + filename

            return Response({"status": True,"message": "Excel report generated successfully.","data": file_url})        
        except Exception as e:
            return Response({"status": False,'message': 'Something went wrong','error': str(e)})
        
from django.utils.text import Truncator
from django.utils.timezone import localtime
class DocumentPDFGenerateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['document_title', 'document_number', 'document_description', 'document_type__name']
    ordering_fields = ['document_title', 'created_at']
    
    queryset = Document.objects.all()

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user 
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            document_current_status = request.query_params.get('document_current_status', None)

            # Admin or Doc Admin: view all or by department
            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                if department_id:
                    queryset = Document.objects.filter(user__department_id=department_id).order_by('-id')
                else:
                    queryset = Document.objects.all().order_by('-id')

            # Reviewer: can only view documents visible to them
            elif user.groups.filter(name='Reviewer').exists():
                queryset = Document.objects.filter(visible_to_users=user).order_by('-id')

            # Department-based filtering for other users
            elif user_department:
                queryset = Document.objects.filter(user__department=user_department).order_by('-id')
            else:
                queryset = Document.objects.none()

            # Further filtering by document_current_status if provided
            if document_current_status:
                queryset = queryset.filter(document_current_status=document_current_status)

            # Apply any filters set in the queryset
            queryset = self.filter_queryset(queryset)

            documents_data = queryset.select_related(
                'document_type', 
                'document_current_status', 
                'select_template', 
                'assigned_to',
            ).values(
                'document_title', 
                'document_number', 
                'parent_document__id',
                'document_type__document_name',  
                'document_current_status__status',
                'select_template__template_name',
                'assigned_to__first_name',
                'assigned_to__last_name',
                'version', 
                'created_at',
            )
            max_lengths = {
            'document_title': max(len(doc['document_title'] or "") for doc in documents_data),
            'document_number': max(len(doc['document_number'] or "") for doc in documents_data),
            'document_type': max(len(doc['document_type__document_name'] or "") for doc in documents_data),
            'assigned_to': max(len(f"{doc['assigned_to__first_name']} {doc['assigned_to__last_name']}" or "") for doc in documents_data),
            'version': max(len(doc['version'] or "") for doc in documents_data),
            'created_at': max(len(str(localtime(doc['created_at']))) for doc in documents_data),
            }

            # Context for rendering HTML template
            context = {
                'documents': documents_data,
                'max_lengths': max_lengths,
                'current_date': timezone.now().strftime('%d-%m-%Y')
            }

            # Load the HTML template
            html_template_path = settings.BASE_DIR / 'templates' / 'document_list_pdf_report.html'
            template = get_template(html_template_path)

            # Render HTML with context data
            html = template.render(context)

            # Create a BytesIO buffer to hold the PDF data in memory
            buffer = BytesIO()

            # Use xhtml2pdf to generate the PDF from the HTML
            pisa_status = pisa.CreatePDF(html, dest=buffer)

            # Check if there were any errors during PDF generation
            if pisa_status.err:
                return Response({
                    'status': False,
                    'message': 'Error occurred while generating the PDF.',
                    'data': {}
                })

            # Get timestamp for PDF file name
            timestamp = int(time.time())
            file_name = f"document_list_{timestamp}.pdf"

            # Define the file path to save the PDF in 'document_certificare' folder
            file_path = os.path.join(settings.MEDIA_ROOT, 'document_certificare', file_name)

            # Create the folder if it doesn't exist
            folder_path = os.path.dirname(file_path)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)

            # Save the generated PDF to the file system
            with open(file_path, 'wb') as f:
                f.write(buffer.getvalue())

            # Build the file URL
            pdf_file_url = f"{settings.MEDIA_URL}document_certificare/{file_name}"
            full_pdf_file_url = f"{request.scheme}://{request.get_host()}{pdf_file_url}"

            # Return response with the PDF file URL
            return Response({
                'status': True,
                'message': 'PDF file generated and saved successfully.',
                'data': {'file_url': full_pdf_file_url}
            })

        except Exception as e:
            return Response({
                'status': False,
                'message': str(e),
                'data': {}
            })
        

class GetObsoleteStatusDataToDocAdminUserOnly(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all()

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user

            if user.groups.filter(name='Doc Admin').exists():
                queryset = Document.objects.filter(document_current_status = 12).order_by('-id')
                serializer = DocumentviewSerializer(queryset, many=True,context={'request': request})
                data = serializer.data
                return Response({"status": True, "message": "Documents retrieved successfully", "data": data})
            else:
                return Response({"status": False, "message": "You are not authorized to view this data"})
        except Exception as e:
            return Response({"status": False, "message": 'Something went wrong', 'error': str(e)})
        
class DocumentDeleteViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentSerializer
    queryset = Document.objects.all()

    def destroy(self, request, *args, **kwargs):
        try:
            document_id = request.data.get('document_id')

            if not Document.objects.filter(id=document_id).exists():
                return Response({"status": False, "message": "Document ID not found"})

            # Retrieve and delete the document
            document_object = Document.objects.get(id=document_id)
            document_object.delete()

            return Response({"status": True, "message": "Document deleted successfully"})

        except Exception as e:
            return Response({"status": False, 'message': 'Something went wrong', 'error': str(e)})
        
from docx import Document as D

# Assuming you have your Document model and DocumentdataSerializer defined elsewhere
# from your_app.models import Document
# from your_app.serializers import DocumentdataSerializer

class DocumentTemplateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentdataSerializer
    queryset = Document.objects.all()
    lookup_field = 'document_id'

    def replace_text_in_paragraphs(self, paragraphs, replacements):
        """Replaces placeholders inside paragraphs, handling multi-run text formatting."""
        for para in paragraphs:
            full_text = "".join(run.text for run in para.runs)  # Combine all runs
            
            modified = False
            for key, value in replacements.items():
                placeholder = f"{{{{ {key} }}}}"  # Ensure format matches docx
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, value)  # Replace placeholders
                    modified = True

            if modified:
                para.clear()  # Clear old content
                para.add_run(full_text)  # Add modified text back

    def replace_text_in_tables(self, tables, replacements):
        """Replaces placeholders inside table cells, ensuring full cell replacement."""
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text = cell.text  # Get entire cell content

                    modified = False
                    for key, value in replacements.items():
                        placeholder = f"{{{{ {key} }}}}"
                        if placeholder in full_text:
                            full_text = full_text.replace(placeholder, value)
                            modified = True

                    if modified:
                        cell.text = full_text  # Replace text in cell

    def list(self, request, *args, **kwargs):
        document_id = self.kwargs.get('document_id')

        if document_id is None:
            return Response({"status": False, "message": "document_id parameter is required"})

        try:
            document = Document.objects.get(id=document_id)
            serializer = self.get_serializer(document, context={'request': request})
            template_url = serializer.data.get('template_url')
            latest_comment = NewDocumentCommentsData.objects.filter(document=document).first()
            front_file_url = latest_comment.front_file_url.url if latest_comment and latest_comment.front_file_url else None
           
            if not template_url:
                return Response({"status": False, "message": "Template URL not found in document data."})

            replacements = {
                "document_title": document.document_title,
                "document_number": document.document_number,
                "version": document.version,
                "department": document.user.department.department_name if document.user.department else 'UnknownDepartment',
            }

            template_path = os.path.join(settings.MEDIA_ROOT, 'templates', f'template_{document_id}.docx')
            output_dir = os.path.join(settings.MEDIA_ROOT, 'generated_docs')
            os.makedirs(os.path.dirname(template_path), exist_ok=True)
            os.makedirs(output_dir, exist_ok=True)

            response = requests.get(template_url)
            if response.status_code != 200:
                return Response({"status": False, "message": "Failed to download template from URL"})

            with open(template_path, 'wb') as f:
                f.write(response.content)

            doc = D(template_path)

            self.replace_text_in_paragraphs(doc.paragraphs, replacements)
            self.replace_text_in_tables(doc.tables, replacements)

            timestamp = int(time.time())
            output_filename = f"training_document_{timestamp}.docx"
            output_path = os.path.join(output_dir, output_filename)
            doc.save(output_path)

            document.generatefile = output_filename
            document.save()

            file_url = f"{settings.MEDIA_URL}generated_docs/{output_filename}"
            full_file_url = f"{request.scheme}://{request.get_host()}{file_url}"
            

            return Response({
                "status": True,
                "message": "Template processed successfully",
                "data": front_file_url if front_file_url else template_url,
            })

        except Document.DoesNotExist:
            return Response({"status": False, "message": "Document not found"})
        except Exception as e:
            return Response({"status": False, "message": str(e)})


class TemplateCreateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = TemplateModel.objects.all()
    
    def create(self, request, *args, **kwargs):
        try:
            user = self.request.user
            template_name = request.data.get('template_name')
            template_doc = request.FILES.get('template_doc')  # Get the uploaded file
            
            if not template_name:
                return Response({"status": False, "message": "Template name is required", "data": []})
            if not template_doc:
                return Response({"status": False, "message": "Template document is required", "data": []})
            
            template = TemplateModel.objects.create(
                template_name=template_name,
                user=user, 
                template_doc=template_doc
            )
            
            return Response({"status": True, "message": "Template created successfully"})
        
        except Exception as e:
            return Response({"status": False, "message": str(e), "data": []})
        
class TemplateViewSet(viewsets.ReadOnlyModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = TemplateSerializer
    queryset = TemplateModel.objects.all().order_by('-id')
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['template_name', 'user__username']
    ordering_fields = ['template_name', 'created_at']

    def list(self, request, *args, **kwargs):
        try:
            queryset = self.filter_queryset(self.get_queryset())
            if queryset.exists():
                serializer = self.get_serializer(queryset, many=True)
                return Response({
                    "status": True,
                    "message": "Templates fetched successfully",
                    'total': queryset.count(),
                    'data': serializer.data
                })
            else:
                return Response({
                    "status": True,
                    "message": "No templates found",
                    "total": 0,
                    "data": []
                })
        except Exception as e:
            return Response({
                "status": False,
                'message': 'Something went wrong',
                'error': str(e)
            })

class TemplateUpdateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = TemplateModel.objects.all()
    lookup_field = 'temp_id'

    def update(self, request, *args, **kwargs):
        try:
            template_id = self.kwargs.get('temp_id') 

            try:
                template = TemplateModel.objects.get(id=template_id)
            except TemplateModel.DoesNotExist:
                return Response({'status': False, 'message': 'Template not found'})

            template_name = request.data.get('template_name')
            template_doc = request.FILES.get('template_doc') 

            if template_name is not None:
                template.template_name = template_name
            if template_doc is not None:
                template.template_doc = template_doc  

            template.save()

            return Response({"status": True, "message": "Template updated successfully"})
        
        except Exception as e:
            return Response({"status": False, "message": 'Something went wrong', 'error': str(e)})
        

class TemplateDocumentViewSet(viewsets.ViewSet):
    permission_classes = [permissions.IsAuthenticated]
    lookup_field = 'template_id'


    def list(self, request, *args, **kwargs):
        template_id = self.kwargs.get('template_id')

        if not template_id:
            return Response({"status": False, "message": "template_id parameter is required"})

        try:
            template = TemplateModel.objects.get(id=template_id)
            serializer = TemplateDocumentSerializer(template, context={'request': request})
            return Response({
                "status": True,
                "message": "Template document fetched successfully",
                "data": serializer.data
            })
        except TemplateModel.DoesNotExist:
            return Response({"status": False, "message": "Template not found"})


        
class DynamicStatusCreateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DynamicStatus.objects.all()
    serializer_class = DynamicStatusSerializer

    def create(self, request, *args, **kwargs):
        try:
            user = request.user
            status_value = request.data.get('status')

          
            if not status_value:
                return Response({"status": False, "message": "Status is required"})

            dynamic_status = DynamicStatus.objects.create(
                user=user,
                status=status_value
            )

            return Response({"status": True, "message": "Dynamic status created successfully"})

        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})
        

class DynamicStatusListViewSet(viewsets.ReadOnlyModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DynamicStatus.objects.all().order_by('-id')
    serializer_class = DynamicStatusSerializer
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['status', 'user__username']
    ordering_fields = ['status', 'created_at']

    def list(self, request, *args, **kwargs):
        try:
            queryset = self.filter_queryset(self.get_queryset())

            if queryset.exists():
                serializer = self.get_serializer(queryset, many=True)
                return Response({
                    "status": True,
                    "message": "Dynamic statuses fetched successfully",
                    'total': queryset.count(),
                    'data': serializer.data
                })
            else:
                return Response({
                    "status": True,
                    "message": "No dynamic statuses found",
                    "total": 0,
                    "data": []
                })
        except Exception as e:
            return Response({
                "status": False, 
                'message': 'Something went wrong', 
                'error': str(e)
            })
        

class DynamicStatusUpdateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DynamicStatusSerializer
    queryset = DynamicStatus.objects.all()
    lookup_field = 'dynamic_status_id'  

    def update(self, request, *args, **kwargs):
        try:
            dynamic_status_id = self.kwargs.get('dynamic_status_id')
            dynamic_status = DynamicStatus.objects.get(id=dynamic_status_id)

            status_value = request.data.get('status')

          
            if status_value is not None:
                dynamic_status.status = status_value
            
            dynamic_status.save()
            return Response({"status": True, "message": "Dynamic status updated successfully"})

        except DynamicStatus.DoesNotExist:
            return Response({"status": False, "message": "Dynamic status not found"})
        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})

# class DynamicStatusDeleteViewSet(viewsets.ModelViewSet):
#     permission_classes = [permissions.IsAuthenticated]
#     queryset = DynamicStatus.objects.all()
#     serializer_class = DynamicStatusSerializer
#     lookup_field = 'dynamic_status_id'  

#     def destroy(self, request, *args, **kwargs):
#         try:
#             dynamic_status_id = self.kwargs.get('dynamic_status_id')
#             dynamic_status = DynamicStatus.objects.get(id=dynamic_status_id)      
#             dynamic_status.delete()
#             return Response({"status": True, "message": "Dynamic status deleted successfully"})

#         except DynamicStatus.DoesNotExist:
#             return Response({"status": False, "message": "Dynamic status not found"})
#         except Exception as e:
#             return Response({"status": False, "message": "Something went wrong", "error": str(e)})

class DocumentDetailsCreateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DocumentDetails.objects.all()
    serializer_class = DocumentDetailsSerializer
    
    def create(self, request, *args, **kwargs):
        try:
            user = self.request.user
            document_id = request.data.get('document_id') 
            document_data = request.data.get('document_data') 

            # Check if the required fields are provided
            if not document_id:
                return Response({"status": False, "message": "Document ID is required", "data": []})
            if not document_data:
                return Response({"status": False, "message": "Document data is required", "data": []})

            # Fetch the document instance based on the provided document_id
            try:
                document = Document.objects.get(id=document_id)
            except Document.DoesNotExist:
                return Response({"status": False, "message": "Document not found", "data": []})

            # Create a new DocumentDetails entry
            document_details = DocumentDetails.objects.create(
                user=user,
                document=document, 
                document_data=document_data
            )
            
            return Response({"status": True, "message": "Document created successfully"})
        
        except Exception as e:
            return Response({"status": False, "message": str(e), "data": []})


class DocumentDetailsUpdateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DocumentDetails.objects.all()
    lookup_field = 'docdetail_id'  
    serializer_class = DocumentDetailsSerializer

    def update(self, request, *args, **kwargs):
        try:
            documentdetail_id = self.kwargs.get('docdetail_id')

            try:
                document_details = DocumentDetails.objects.get(id=documentdetail_id)
            except DocumentDetails.DoesNotExist:
                return Response({"status": False, "message": "Document not found"})

            # Get new data from the request
            document_data = request.data.get('document_data')
            document_id = request.data.get('document_id') 


            if document_data is not None:
                document_details.document_data = document_data
            
            # if document_id is not None:
            #     document_details.document_id = document_id
            version_number = document_details.document.version  # Get the current version
            new_version = increment_version(version_number)
            document_details.document.version = new_version
            document_details.save() 

            return Response({
                "status": True,
                "message": "Document Details updated successfully",
            })

        except Exception as e:
            return Response({
                "status": False,
                "message": "Something went wrong",
                "error": str(e)
            })
        

class DocumentDetailsViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentDetailsSerializer
    queryset = DocumentDetails.objects.all().order_by('-id')
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['document__name', 'user__username'] 
    ordering_fields = ['created_at']

    def list(self, request):
        queryset = self.filter_queryset(self.get_queryset()) 
        
        try:
            if queryset.exists():
                serializer = DocumentDetailsSerializer(queryset, many=True)
                return Response({
                    "status": True,
                    "message": "Documents fetched successfully",
                    'total': queryset.count(),
                    'data': serializer.data
                })
            else:
                return Response({
                    "status": True,
                    "message": "No Documents found",
                    "total": 0,
                    "data": []
                })
        except Exception as e:
            return Response({"status": False, 'message': 'Something went wrong', 'error': str(e)})


class DocumentApproveActionCreateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DocumentAuthorApproveAction.objects.all()

    def create(self, request, *args, **kwargs):
        try:
            user = self.request.user
            document = request.data.get('document_id')
            status_id = request.data.get('status')
            remark = request.data.get('remark')
            visible_to_users = request.data.get('visible_to_users', [])
            approver = request.data.get('approver')
            doc_admin = request.data.get('doc_admin')

            # Ensure required fields are provided
            if not document:
                return Response({"status": False, "message": "Document are required"})
            if not status_id:
                return Response({"status": False, "message": "Status is required"})
            if not remark:
                return Response({"status": False, "message": "remark is required"})

             # Fetch related document and status objects
            try:
                document = Document.objects.get(id=document)
            except Document.DoesNotExist:
                return Response({"status": False, "message": "Document not found"})

            try:
                status = DynamicStatus.objects.get(id=status_id)
            except DynamicStatus.DoesNotExist:
                return Response({"status": False, "message": "Invalid status ID"})
            
            status_underreivew = DynamicStatus.objects.get(id=3)
            # DocumentReviewerAction.objects.filter(document=document).delete()
            DocumentReviewerAction.objects.filter(document=document).update(status_approve=status_underreivew)
            SendBackofUser.objects.filter(document=document).update(is_send_back=False)
            ReviewByUser.objects.filter(document=document).update(is_reviewed=False)
            UserWiseSendBackView.objects.filter(document=document).update(is_done=False)

            # Update visible_to_users
            if isinstance(visible_to_users, str):
                import json
                try:
                    visible_to_users = json.loads(visible_to_users)
                except json.JSONDecodeError:
                    return Response({
                        "status": False,
                        "message": "Invalid format for visible_to_users. Provide a valid list of IDs.",
                    })
            if visible_to_users:
                document.visible_to_users.set(visible_to_users)
                
             # Update approver
            if approver:
                try:
                    approver_user = CustomUser.objects.get(id=approver)
                    document.approver = approver_user
                except CustomUser.DoesNotExist:
                    return Response({"status": False, "message": "Approver user not found"})

            # Update doc_admin
            if doc_admin:
                try:
                    doc_admin_user = CustomUser.objects.get(id=doc_admin)
                    document.doc_admin = doc_admin_user
                except CustomUser.DoesNotExist:
                    return Response({"status": False, "message": "Doc Admin user not found"})

            # Save updated document
            document.save()
           
            document_approve_action = DocumentAuthorApproveAction.objects.create(
                user=user,
                document=document,
                status_approve=status,
                remarks_author = remark,
            )
            approve_action = DocApprove.objects.create(
                user=user,
                document=document,
                status_approve=status
            )
            document.author = user
            document.document_current_status = status
            document.form_status = None
            document.assigned_to = None
            version_number = document.version
            new_version = increment_version(version_number)
            document.version = new_version
            document.save()
            
            Archived.objects.create(
                document=document,
                version=version_number
            )
            UserWiseSendBackView.objects.filter(document=document).update(is_done=False)
            
            document_title = document.document_title
            reviewer_group = Group.objects.get(name='Reviewer')
            reviewers = CustomUser.objects.filter(groups=reviewer_group)
            department_users = CustomUser.objects.filter(department=user.department)
            # users_to_notify = reviewers.union(department_users).distinct()
            # Combine the querysets and eliminate duplicates
            users_to_notify = CustomUser.objects.filter(
                id__in=set(reviewers.values_list('id', flat=True)) | 
                set(department_users.values_list('id', flat=True))
            )
            send_document_update_email(user, document_title, users_to_notify)
            
            return Response({"status": True, "message": "Document approval action created successfully"})

        except DocumentDetails.DoesNotExist:
            return Response({"status": False, "message": "Invalid document details ID"})
        except DynamicStatus.DoesNotExist:
            return Response({"status": False, "message": "Invalid status ID"})
        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})


class DocumentReviewerActionCreateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DocumentReviewerAction.objects.all()

    def create(self, request, *args, **kwargs):
        try:
            user = self.request.user
            document_id = request.data.get('document_id')
            status_id = request.data.get('status')
            remark = request.data.get('remark')

            # Validate required fields
            if not document_id:
                return Response({"status": False, "message": "Document is required"})
            if not status_id:
                return Response({"status": False, "message": "Status is required"})
            if not remark:
                return Response({"status": False, "message": "Remark is required"})

            # Fetch the document
            try:
                document = Document.objects.get(id=document_id)
            except Document.DoesNotExist:
                return Response({"status": False, "message": "Invalid Document ID"})

            # Fetch the status
            try:
                status = DynamicStatus.objects.get(id=status_id)
            except DynamicStatus.DoesNotExist:
                return Response({"status": False, "message": "Invalid Status ID"})

            review_done, created = ReviewByUser.objects.get_or_create(
                user=user,
                document=document,
                defaults={"is_reviewed": True}  # Set is_reviewed=True if new
            )
            if not created:
                review_done.is_reviewed = True
                review_done.save()
            document_reviewer_action = DocumentReviewerAction.objects.create(
                user=user,
                document=document,
                status_approve=status,
                remarks_reviewer=remark
            )
            approve_action = DocApprove.objects.create(
                user=user,
                document=document,
                status_approve=status
            )

            # Fetch the status for 'Send Back' and 'Approved'
            send_back_status = DynamicStatus.objects.get(id=8)
            approve_status = DynamicStatus.objects.get(id=4)

            # Get the total number of assigned reviewers
            total_reviewers = document.visible_to_users.count()

            # Check how many reviewers have approved and how many sent it back
            approved_reviews = DocumentReviewerAction.objects.filter(
                document=document,
                status_approve=status
            ).count()
            
            sent_back_reviews = DocumentReviewerAction.objects.filter(
                document=document,
                status_approve=send_back_status
            ).count()
            print(sent_back_reviews,"sent_back_reviews")
            print(approved_reviews,"approved_reviews")
            print(total_reviewers,"total_")
            # If any reviewer sends it back, don't change the status
            if sent_back_reviews > 0:
                document.document_current_status = send_back_status  # 'Send Back'
            elif approved_reviews == total_reviewers:
                # Only change to 'Approved' if all reviewers have approved
                document.document_current_status = approve_status  # 'Approved'
            # Otherwise, don't change the status
            
            # Proceed with the rest of the document update
            document.assigned_to = None
            version_number = document.version  # Get the current version
            new_version = increment_version(version_number)
            document.version = new_version
            document.save()

            Archived.objects.create(
                document=document,
                version=version_number
            )

            # Send approval notification to approvers and department users
            document_title = document.document_title
            approver_user = Group.objects.get(name='Approver')
            approvers = CustomUser.objects.filter(groups=approver_user)
            department_users = CustomUser.objects.filter(department=document.user.department)

            # Use union without distinct() to avoid the error
            users_to_notify = approvers.union(department_users)

            send_document_approval_email(user, document_title, users_to_notify)

            return Response({
                "status": True,
                "message": "Document reviewer action created successfully",
            })

        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})

class GetSendBackActionViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = SendBackofUser.objects.all()
    
    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            document_id = request.query_params.get('document_id')
            if not document_id:
                return Response({"status": False, "message": "Document ID is required"})
            send_back_requests = SendBackofUser.objects.filter(user=user, document_id=document_id)
            serialized_data = SendBackofUserSerializer(send_back_requests, many=True)
            return Response({
                "status": True,
                "message": "Send back requests retrieved successfully",
                "send_back_requests": serialized_data.data
            })
        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})

class DocumentApproverActionCreateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DocumentApproverAction.objects.all()

    def create(self, request, *args, **kwargs):
        try:
            user = self.request.user
            document_id = request.data.get('document_id')
            status_id = request.data.get('status')
            remark = request.data.get('remark')


            # Validate required fields
            if not document_id:
                return Response({"status": False, "message": "Document is required"})
            if not status_id:
                return Response({"status": False, "message": "Status is required"})
            if not remark:
                return Response({"status": False, "message": "remark is required"})
            # Fetch the document
            try:
                document = Document.objects.get(id=document_id)
            except Document.DoesNotExist:
                return Response({"status": False, "message": "Invalid Document ID"})

            # Fetch the status
            try:
                status = DynamicStatus.objects.get(id=status_id)
            except DynamicStatus.DoesNotExist:
                return Response({"status": False, "message": "Invalid Status ID"})

           
            # Create the reviewer action
            document_approver_action = DocumentApproverAction.objects.create(
                user=user,
                document=document,
                status_approve=status,
                remarks_approver = remark
            )
            approve_action = DocApprove.objects.create(
                user=user,
                document=document,
                status_approve=status
            )


            # Update document current status only if all reviewers have approved
            # if all_reviewers_approved:
            document.document_current_status = status
            document.assigned_to = None
            version_number = document.version  # Get the current version
            new_version = get_new_version(version_number)
            document.version = new_version
            document.save()
            
            Archived.objects.create(
                document=document,
                version=version_number
            )
            # if new_version.endswith(".0"):  
            #     previous_major_version = str(int(new_version.split(".")[0]) - 1) + ".0"

            #     old_doc = Document.objects.filter(
            #         document_number=document.document_number,
            #         version=previous_major_version
            #     ).first()

            #     if old_doc:
            #         old_doc.document_current_status = DynamicStatus.objects.get(id=15)
            #         old_doc.save()
            
            # Send approval notification to approvers and department users
            document_title = document.document_title
            approver_user = Group.objects.get(name='Doc Admin')
            approvers = CustomUser.objects.filter(groups=approver_user)
            department_users = CustomUser.objects.filter(department=document.user.department)

            # Use union without distinct() to avoid the error
            users_to_notify = approvers.union(department_users)

            send_document_approval_email(user, document_title, users_to_notify)

            return Response({
                "status": True,
                "message": "Document approver action created successfully",
            })

        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})


class DocumentDocAdminActionCreateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DocumentDocAdminAction.objects.all()

    def create(self, request, *args, **kwargs):
        try:
            user = self.request.user
            document_id = request.data.get('document_id')
            status_id = request.data.get('status')


            # Validate required fields
            if not document_id:
                return Response({"status": False, "message": "Document is required"})
            if not status_id:
                return Response({"status": False, "message": "Status is required"})


            # Fetch the document
            try:
                document = Document.objects.get(id=document_id)
            except Document.DoesNotExist:
                return Response({"status": False, "message": "Invalid Document ID", "data": []})

            # Fetch the status
            try:
                status = DynamicStatus.objects.get(id=status_id)
            except DynamicStatus.DoesNotExist:
                return Response({"status": False, "message": "Invalid Status ID", "data": []})

           
            # Create the reviewer action
            document_docAdmin_action = DocumentDocAdminAction.objects.create(
                user=user,
                document=document,
                status_approve=status,
            )

            approve_action = DocApprove.objects.create(
                user=user,
                document=document,
                status_approve=status
            )

            document.document_current_status = status
            document.save()

            # Log actions based on status
            if status_id == 7:  
                DocumentEffectiveAction.objects.create(
                    user=user,
                    documentdetails_effective=document,
                    remarks_effective=status,

                )
                user_department = document.user.department
                department_users = CustomUser.objects.filter(department=user_department)
                for department_user in department_users:
                    send_document_doc_admin_effective_email(department_user, document, status)

            elif status_id == 6:  
                DocumentReleaseAction.objects.create(
                    user=user,
                    documentdetails_release=document,
                    status_release=status,
                )
                user_department = document.user.department
                department_users = CustomUser.objects.filter(department=user_department)
                for department_user in department_users:
                    send_document_doc_admin_release_email(department_user, document, status)
            return Response({"status": True,"message": "Document Doc Admin action created successfully", "data": []})

        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})

class DocumentIdwiseAuthorReviewerApproverDocAdminViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = Document.objects.all()
    
    def list(self, request,*args, **kwargs):
        try:
            document_id = self.kwargs.get('document_id')
            if not document_id:
                return Response({"status": False, "message": "Document ID is required"})
            try:
                document_instance = Document.objects.get(id=document_id)
                author = document_instance.author
                reviewer = document_instance.visible_to_users
                doc_admin = document_instance.doc_admin
                approver = document_instance.approver
                data = {
                    "author": author.username,
                    "reviewer": reviewer.values_list('username', flat=True),
                    "doc_admin": doc_admin.username,
                    "approver": approver.username
                }
                return Response({"status": True, "message": "Document ID wise author, reviewer, doc admin, approver fetched successfully", "data": data})
            except Document.DoesNotExist:
                return Response({"status": False, "message": "Invalid Document ID", "data": []})
        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})
            

class DocumentSendBackActionCreateViewSet(viewsets.ViewSet):
    permission_classes = [permissions.IsAuthenticated]

    def create(self, request, *args, **kwargs):
        try:
            user = request.user
            document_id = request.data.get('document_id')
            assigned_to_id = request.data.get('assigned_to')
            status_id = request.data.get('status_sendback')
            group_id = request.data.get('assign_user_group')
            remark = request.data.get('remark')


            # Validate required fields
            if not document_id or not assigned_to_id or not status_id:
                return Response({
                    "status": False,
                    "message": "Document ID, Assigned User ID, and Status ID are required"
                })
                
            if not remark:
                return Response({"status": False, "message": "remark is required"})


            # Fetch the document
            try:
                document = Document.objects.get(id=document_id)
            except Document.DoesNotExist:
                return Response({"status": False, "message": "Invalid Document ID"})

            # Fetch the assigned user
            try:
                assigned_to = CustomUser.objects.get(id=assigned_to_id)
            except CustomUser.DoesNotExist:
                return Response({"status": False, "message": "Invalid Assigned User ID"})

            # Fetch the status
            try:
                status = DynamicStatus.objects.get(id=status_id)
            except DynamicStatus.DoesNotExist:
                return Response({"status": False, "message": "Invalid Status ID"})
            
            if status.id == 8:
                send_back, created = SendBackofUser.objects.get_or_create(
                user=user, 
                document=document,
                defaults={"is_send_back": True} 
                )
                if not created: 
                    send_back.is_send_back = True
                    send_back.save()

            # Log the send-back action
            send_back_action = DocumentSendBackAction.objects.create(
                user=user,
                document=document,
                status_sendback=status,
                group = group_id,
                remarks_sendback = remark
            )
            send_for_user, created = UserWiseSendBackView.objects.get_or_create(
                user=user,
                document=document,
                defaults={"is_done": True}
            )
            if not created:
                send_for_user.is_done = True
                send_for_user.save()

            # Update the document's assigned user and reason
            document.assigned_to = assigned_to
            document.assigned_to_group = group_id
            document.document_current_status = status
            document.save()
            document_title = document.document_title
            send_document_sendback_email(assigned_to, document_title)
            
            return Response({
                "status": True,
                "message": "Document sent back successfully",
            })

        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})

from django.utils.timezone import make_aware

class DocumentStatusHandleViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DocumentReleaseAction.objects.all()

    def create(self, request, *args, **kwargs):
        try:
            user = self.request.user
            document_id = request.data.get('document_id')
            status_id = request.data.get('status_id')
            effective_date = request.data.get('effective_date')
            revision_date = request.data.get('revision_date')


            if not document_id:
                return Response({"status": False, "message": "Document details are required"})
            # if not status_id:
            #     return Response({"status": False, "message": "Status is required"})
            
            
            try:
                status_id = int(status_id)
            except (TypeError, ValueError):
                return Response({"status": False, "message": f"Invalid status ID: {status_id}"})

            status_release = DynamicStatus.objects.get(id=status_id)

            try:
                document = Document.objects.get(id=document_id)
            except Document.DoesNotExist:
                return Response({"status": False, "message": "Invalid Document ID"})
            
             # Convert effective_date and revision_date to timezone-aware datetime objects
            if effective_date:
                try:
                    naive_effective_date = datetime.strptime(effective_date, "%d-%m-%Y")
                    effective_date = make_aware(naive_effective_date)  # Convert to timezone-aware datetime
                except ValueError:
                    return Response({"status": False, "message": "Invalid effective date format. Expected format: dd-mm-yyyy"})
            if revision_date:
                try:
                    naive_revision_date = datetime.strptime(revision_date, "%d-%m-%Y")
                    revision_date = make_aware(naive_revision_date)  # Convert to timezone-aware datetime
                except ValueError:
                    return Response({"status": False, "message": "Invalid revision date format. Expected format: dd-mm-yyyy"})
            
            # Determine which model to use based on status_id
            if status_id == 6:
                document_release_action = DocumentReleaseAction.objects.create(
                    user=user,
                    document_id=document_id,
                    status_release=status_release,
                )
                document.document_current_status = status_release
                document.save()
                # send email to the all user who's department is same as the who created the document
                user_department = document.user.department
                department_users = CustomUser.objects.filter(department=user_department)
                for department_user in department_users:
                    send_document_release_email(department_user, document, status_release)
                return Response({"status": True, "message": "Document release action created successfully"})
            elif status_id == 7:
                document_effective_action = DocumentEffectiveAction.objects.create(
                    user=user,
                    document_id=document_id,
                    status_effective=status_release,
                    effective_date = effective_date,
                )
                # Update the document's current status
                document.effective_date = effective_date
                document.revision_date = revision_date
                document.document_current_status = status_release
                version_number = document.version  # Get the current version
                new_version = increment_version(version_number)
                document.version = new_version 
                document.save()
                
                Archived.objects.create(
                document=document,
                version=version_number
                )
                
                user_department = document.user.department
                department_users = CustomUser.objects.filter(department=user_department)
                for department_user in department_users:
                    send_document_effective_email(department_user, document, status_release)
                return Response({"status": True, "message": "Document effective action created successfully"})
            else:
                return Response({"status": False, "message": "Invalid status ID"})            

        except Document.DoesNotExist:
            return Response({"status": False, "message": "Invalid document ID"})
        except DynamicStatus.DoesNotExist:
            return Response({"status": False, "message": "Invalid status ID"})
        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})


# class DocumentEffectiveActionCreateViewSet(viewsets.ModelViewSet):
#     permission_classes = [permissions.IsAuthenticated]
#     queryset = DocumentEffectiveAction.objects.all()

#     def create(self, request, *args, **kwargs):
#         try:
#             user = self.request.user
#             document_id = request.data.get('documentdetails_effective')
#             status_id = request.data.get('status_effective')

#             if not document_id:
#                 return Response({"status": False, "message": "Document details are required"})
#             if not status_id:
#                 return Response({"status": False, "message": "Status is required"})

#             documentdetails_effective = DocumentDetails.objects.get(id=document_id)
#             status_effective = DynamicStatus.objects.get(id=status_id)

#             document_effective_action = DocumentEffectiveAction.objects.create(
#                 user=user,
#                 documentdetails_effective=documentdetails_effective,
#                 status_effective=status_effective
#             )

#             return Response({"status": True, "message": "Document effective action created successfully"})

#         except DocumentDetails.DoesNotExist:
#             return Response({"status": False, "message": "Invalid document details ID"})
#         except DynamicStatus.DoesNotExist:
#             return Response({"status": False, "message": "Invalid status ID"})
#         except Exception as e:
#             return Response({"status": False, "message": "Something went wrong", "error": str(e)})
        
#     def list(self, request, *args, **kwargs):
#         try:
#             user = self.request.user

#             queryset = DocumentEffectiveAction.objects.filter(user=user)

#             if queryset.exists():
#                 serializer = DocumentEffectiveActionSerializer(queryset, many=True)
#                 data = serializer.data
#                 return Response({"status": True, "message": "Document effective actions fetched successfully", "data": data})
#             else:
#                 return Response({"status": False, "message": "No document effective actions found", "data": []})
#         except Exception as e:
#             return Response({"status": False, "message": "Something went wrong", "error": str(e)})


         
class DocumentReviseActionViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DocumentRevisionAction.objects.all()

    def create(self, request, *args, **kwargs):
        try:
            user = self.request.user
            document_id = request.data.get('document_id',None)
            status_id = request.data.get('status_id',None)
            request_action_id = request.data.get('request_action_id',None)
            action_status = request.data.get('action_status',None)
            print(status_id, "status_id")
            if not document_id:
                return Response({"status": False, "message": "Document ID is required"})
            if not status_id:
                return Response({"status": False, "message": "Status ID is required"})
            if not request_action_id:
                return Response({"status": False, "message": "Revision request ID is required"})
            if not action_status or action_status not in ['approved', 'rejected']:
                return Response({"status": False, "message": "Action status is required and must be 'approved' or 'rejected'"})
            
            document = Document.objects.get(id=document_id)
            status_revision = DynamicStatus.objects.get(id=status_id)
            revision_request = DocumentRevisionRequestAction.objects.get(id=request_action_id)
            revision_request.status = action_status
            revision_request.save()

            if document.parent_document:
                document_parent = Document.objects.get(id=document.parent_document.id)
                # print(document_parent)
                parent_job_roles = document_parent.job_roles.all()
                document_parent.job_roles.clear()
                    # print(f"Cleared all job roles from parent document {parent.id}")
            # revise_action = DocumentRevisionAction.objects.create(
            #     user=user,
            #     document=document,
            #     status_revision=status_revision
            # )
            # revise_action.save()

            # version_number = document.version
            # new_version = get_new_version(version_number)
            document.is_revised = True
            # document.version = new_version
            # document.document_current_status = status_revision
            document.save()
            message = "Revision request processed successfully."

            if status_id == 10:  # If approved, create a duplicate document entry
                new_version = increment_version(document.version)
                new_document = Document.objects.create(
                    user=user,
                    document_title=document.document_title,
                    parent_document=document.parent_document,
                    workflow = document.workflow,
                    document_operation = document.document_operation,
                    revision_month=document.revision_month,
                    # assigned_to=document.assigned_to,
                    select_template=document.select_template,
                    document_number=document.document_number,
                    document_type=document.document_type,
                    form_status=document.form_status,
                    document_current_status=status_revision,  # Set status to 10 (revise)
                    version=new_version,
                    # is_revised=True,
                    training_required=document.training_required,
                    document_description=document.document_description,
                    # effective_date=document.effective_date,
                    # revision_date=document.revision_date,
                    product_code=document.product_code,
                    equipment_id=document.equipment_id,
                    generatefile=document.generatefile,
                    # author=user
                )
                child_documents = Document.objects.filter(parent_documents=document)
                for child in child_documents:
                    child.parent_document.add(new_document)
                    child.save()
                old_questions = TrainingQuestions.objects.filter(document=document)
                for question in old_questions:
                    TrainingQuestions.objects.create(
                        document=new_document,  # Assign the new document
                        question_type=question.question_type,
                        question_text=question.question_text,
                        options=question.options,
                        correct_answer=question.correct_answer,
                        marks=question.marks,
                        created_by=user,  # Set the user who initiated the revision
                        selected_file_type=question.selected_file_type,
                        selected_file=question.selected_file,
                        question_created_at=timezone.now()
                    )
                message = "Revision request successfully approved, and questions were copied."

            elif status_id == 11:  # If rejected, just update the document status
                document.document_current_status = status_revision  # Set status to 11 (rejected)
                revision_request.is_revise = False # Set status to 11 (rejected)
                document.save()
                message = "Revision request successfully rejected."

            DocumentRevisionAction.objects.create(
                user=user,
                document=document,
                status_revision=status_revision
            )
                # return Response({
                #     "status": True,
                #     "message": "Revise action created successfully",
                # })
            # message = "Revision request successfully " + ("approved" if action_status == "approved" else "rejected")
            return Response({
                "status": True,
                "message": message,
            })
            
                # all_users = CustomUser.objects.filter(department = documentdetails_revise.document.user.department)
                # for user in all_users:
                #     send_document_revise_email(user, documentdetails_revise, status_revise)

        except Document.DoesNotExist:
            return Response({"status": False, "message": "Invalid document ID"})
        except DynamicStatus.DoesNotExist:
            return Response({"status": False, "message": "Invalid status ID"})
        except DocumentRevisionRequestAction.DoesNotExist:
            return Response({"status": False, "message": "Invalid revision request ID"})
        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})


class MasterCopyUserDropdownViewSet(viewsets.ModelViewSet):
    queryset = CustomUser.objects.all().order_by('-id')
    serializer_class = SimpleUserSerializer
    permission_classes = [permissions.IsAuthenticated] 

    def list(self, request, *args, **kwargs):
        try:
            queryset = self.filter_queryset(self.get_queryset())
            serializer = SimpleUserSerializer(queryset, many=True)
            data = serializer.data
            return Response({"status": True, "message": "User list fetched successfully", "data": data})
        except Exception as e:
            return Response({"status": False, "message": str(e), "data": []})
        

class OtherUserDropdownViewSet(viewsets.ModelViewSet):
    queryset = CustomUser.objects.all().order_by('-id')
    serializer_class = SimpleUserSerializer
    permission_classes = [permissions.IsAuthenticated] 

    def list(self, request, *args, **kwargs):
        try:
            queryset = self.filter_queryset(self.get_queryset())
            serializer = SimpleUserSerializer(queryset, many=True)
            data = serializer.data
            return Response({"status": True, "message": "User list fetched successfully", "data": data})
        except Exception as e:
            return Response({"status": False, "message": str(e), "data": []})


class DynamicInventoryCreateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DynamicInventory.objects.all()
    serializer_class = DynamicInventorySerializer

    def create(self, request, *args, **kwargs):
        try:
            inventory_name = request.data.get('inventory_name')
            if not inventory_name:
                return Response({"status": False, "message": "Inventory name is required"})

            dynamic_inventory = DynamicInventory.objects.create(inventory_name=inventory_name)
            serializer = DynamicInventorySerializer(dynamic_inventory)
            return Response({"status": True, "message": "Dynamic inventory created successfully", "data": serializer.data})

        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})

class DynamicInventoryListViewSet(viewsets.ReadOnlyModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DynamicInventory.objects.all().order_by('-id')
    serializer_class = DynamicInventorySerializer
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['inventory_name']
    ordering_fields = ['inventory_name', 'created_at']

    def list(self, request, *args, **kwargs):
        try:
            queryset = self.filter_queryset(self.get_queryset())
            serializer = self.get_serializer(queryset, many=True)
            return Response({
                "status": True,
                "message": "Dynamic inventories fetched successfully",
                'total': queryset.count(),
                'data': serializer.data
            })

        except Exception as e:
            return Response({"status": False, 'message': 'Something went wrong', 'error': str(e)})

class DynamicInventoryUpdateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DynamicInventorySerializer
    queryset = DynamicInventory.objects.all()
    lookup_field = 'inventory_id'

    def update(self, request, *args, **kwargs):
        try:
            # Retrieve the inventory_id from the URL
            inventory_id = self.kwargs.get('inventory_id')
            
            # Retrieve inventory_name from the request data
            inventory_name = request.data.get('inventory_name')
            if not inventory_name:
                raise ValidationError("The field 'inventory_name' is required and cannot be null.")

            # Retrieve the inventory instance and update
            dynamic_inventory = DynamicInventory.objects.get(id=inventory_id)
            dynamic_inventory.inventory_name = inventory_name
            dynamic_inventory.save()

            # Serialize and return the updated data
            serializer = DynamicInventorySerializer(dynamic_inventory)
            return Response({
                "status": True, 
                "message": "Dynamic inventory updated successfully", 
                "data": serializer.data
            })

        except DynamicInventory.DoesNotExist:
            return Response({
                "status": False, 
                "message": "Dynamic inventory not found"
            })

        except ValidationError as ve:
            return Response({
                "status": False, 
                "message": str(ve)
            })

        except Exception as e:
            return Response({
                "status": False, 
                "message": "Something went wrong", 
                "error": str(e)
            })

class DynamicInventoryDeleteViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DynamicInventory.objects.all()
    serializer_class = DynamicInventorySerializer
    lookup_field = 'inventory_id'

    def destroy(self, request, *args, **kwargs):
        try:
            inventory_id = self.kwargs.get('inventory_id')
            dynamic_inventory = DynamicInventory.objects.get(id=inventory_id)
            dynamic_inventory.delete()
            return Response({"status": True, "message": "Dynamic inventory deleted successfully"})

        except DynamicInventory.DoesNotExist:
            return Response({"status": False, "message": "Dynamic inventory not found"})
        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})
        


class DocumentCommentCreateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DocumentComments.objects.all().order_by('-created_at')

    def create(self, request):
        try:
            document_id = request.data.get('document')
            comment_description = request.data.get('comment_description', {})
            
            # Ensure required fields are provided
            if not document_id:
                return Response({'status': False, 'message': 'Document ID is required'})
            if not comment_description:
                return Response({'status': False, 'message': 'Comment description is required'})

            # Create the comment
            comment_obj = DocumentComments.objects.create(
                user=self.request.user,
                document_id=document_id,
                Comment_description=comment_description
            )
            return Response({'status': True, 'message': 'Comment added successfully'})
        except Exception as e:
            return Response({'status': False, 'message': 'Something went wrong', 'error': str(e)})
        
class DocumentCommentsViewSet(viewsets.ModelViewSet):
    queryset = DocumentComments.objects.all().order_by('-created_at')
    serializer_class = DocumentCommentSerializer
    permission_classes = [permissions.IsAuthenticated]
    lookup_field = 'document_id'


    def list(self, request, *args, **kwargs):
        try:
            document_id = self.kwargs.get('document_id')
            queryset = self.filter_queryset(self.get_queryset())

            if document_id:
                queryset = queryset.filter(document_id=document_id)

            serializer = DocumentCommentSerializer(queryset, many=True)
            data = serializer.data
            return Response({"message": "Comment list fetched successfully", "data": data})
        except Exception as e:
            return Response({"message": str(e), "data": []})
        
class DocumentCommentDeleteViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    lookup_field = 'comment_id'

    def destroy(self, request, *args, **kwargs):
        try:
            comment_id = self.kwargs.get('comment_id')
            document_comment = DocumentComments.objects.get(id=comment_id)
            document_comment.delete()
            return Response({"message": "Comment deleted successfully"})

        except DocumentComments.DoesNotExist:
            return Response({"message": "Comment not found"})
        except Exception as e:
            return Response({"message": "Something went wrong", "error": str(e)})
        

class DocumentDetailViewSet(viewsets.ViewSet):
    permission_classes = [permissions.IsAuthenticated]
    lookup_field = 'document_id'
    
    def list(self, request, *args, **kwargs):
        document_id = self.kwargs.get('document_id')

        if not document_id:
            return Response({"status": False, "message": "document_id parameter is required"})

        try:
            document = Document.objects.get(id=document_id)
            serializer = DocumentDetailSerializer(document)
            return Response({
                "status": True,
                "message": "Document details fetched successfully",
                "data": serializer.data
            })
        except Document.DoesNotExist:
            return Response({"status": False, "message": "Document not found"})
        
class DocumentDraftStatusViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = Document.objects.all()

    def update(self, request, *args, **kwargs):
        try:
            user = request.user
            document_id = request.data.get('document_id')
            status_id = request.data.get('status_id')

            # Ensure required fields are provided
            if not document_id:
                return Response({"status": False, "message": "Document ID is required"})
            if not status_id:
                return Response({"status": False, "message": "Status ID is required"})

            # Fetch related document and status objects
            try:
                document = Document.objects.get(id=document_id)
            except Document.DoesNotExist:
                return Response({"status": False, "message": "Invalid Document ID"})

            try:
                status = DynamicStatus.objects.get(id=status_id)
            except DynamicStatus.DoesNotExist:
                return Response({"status": False, "message": "Invalid Status ID"})

            # Update the fields in the Document table
            document.document_current_status = status
            document.form_status = "save_draft"
            document.save()

            return Response({"status": True, "message": "Document Drafted successfully"})

        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})

class DocumentTypeUpdateViewSet(viewsets.ModelViewSet):
    serializer_class = DocumentTypeSerializer
    queryset = DocumentType.objects.all()
    permission_classes = [permissions.IsAuthenticated]
    lookup_field = 'document_type_id'       

    def update(self, request, *args, **kwargs):
        try:
            # Get document_type_id and new document_name from the request
            document_type_id = self.kwargs.get('document_type_id')
            document_name = request.data.get('document_name')

            # Validate document_type_id and document_name
            if not document_type_id:
                return Response({"status": False, "message": "Document type ID is required", "data": []})
            if not document_name:
                return Response({"status": False, "message": "Document name is required", "data": []})

            # Fetch the DocumentType instance
            document_type = DocumentType.objects.filter(id=document_type_id).first()
            if not document_type:
                return Response({"status": False, "message": "Document type not found", "data": []})

            # Update the document_name
            document_type.document_name = document_name
            document_type.save()

            # Serialize the updated instance
            serializer = DocumentTypeSerializer(document_type)
            return Response({"status": True, "message": "Document type updated successfully"})
        except Exception as e:
            return Response({"status": False, "message": str(e), "data": []})


# class DepartmentUsersViewSet(viewsets.ModelViewSet):
#     permission_classes = [permissions.IsAuthenticated]
#     serializer_class = CustomUserdataSerializer

#     def list(self, request, *args, **kwargs):
#         try:
#             # Get the document ID from the request
#             document_id = request.data.get('document_id', None)
#             if not document_id:
#                 return Response({
#                     "status": False,
#                     "message": "Document ID is required.",
#                     "data": []
#                 })

#             # Fetch the document and ensure it exists
#             try:
#                 document = Document.objects.get(id=document_id)
#             except Document.DoesNotExist:
#                 return Response({
#                     "status": False,
#                     "message": "Document not found.",
#                     "data": []
#                 })

#             # Fetch users associated with the document from the DocApprove model
#             approved_users = CustomUser.objects.filter(
#                 docapprove__document=document
#             ).distinct()

#             # Serialize the filtered users
#             serializer = self.serializer_class(approved_users, many=True, context={'request': request})
#             return Response({
#                 "status": True,
#                 "message": "Users fetched successfully.",
#                 "data": serializer.data,
#             })
#         except Exception as e:
#             return Response({
#                 "status": False,
#                 "message": str(e),
#                 "data": [],
#             })


class DepartmentUsersViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    lookup_field = 'document_id'

    def list(self, request, *args, **kwargs):
        try:
            # Get the document ID from the request
            document_id = self.kwargs.get('document_id')
            if not document_id:
                return Response({
                    "status": False,
                    "message": "Document ID is required.",
                    "data": []
                })

            # Fetch the users linked to the document ID in the DocApprove model
            document = Document.objects.filter(id=document_id).first()
            if not document:
                return Response({
                    "status": False,
                    "message": "No users found for the given document ID.",
                    "data": []
                })

            # Prepare data for response
            response_data = []
            if document.author:
                response_data.append({
                    "id": document.author.id,
                    "first_name": f"{document.author.first_name} (Author)",
                    "group_id": [group.id for group in document.author.groups.all()]
                })

            # Return the response
            return Response({
                "status": True,
                "message": "Users fetched successfully.",
                "data": response_data,
            })
        except Exception as e:
            return Response({
                "status": False,
                "message": str(e),
                "data": [],
            })

class PrinterMachines(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]

    def create(self, request, *args, **kwargs):
        try:
            user = self.request.user
            printer_name = request.data.get('printer_name')
            printer_description = request.data.get('printer_location')

            if not printer_description:
                return Response({'status':False,'message':'Print Location is required'})
            
            if not printer_name:
                return Response({'status':False,'message':'Print Name is required'})
            
            printer_obj = PrinterMachinesModel.objects.create(
                user = user,
                printer_name = printer_name,
                printer_description = printer_description
            )
            return Response({'status': True, 'message': 'Printer Added successfully'})
        except Exception as e:
            return Response({'status': False, 'message': 'Something went wrong', 'error': str(e)})
    
    def list(self, request):
        queryset = PrinterMachinesModel.objects.all().order_by('-id')
        
        try:
            if queryset.exists():
                serializer = PrinterSerializer(queryset, many=True)
                return Response({
                    "status": True,
                    "message": "Printer Machine fetched successfully",
                    'total': queryset.count(),
                    'data': serializer.data
                })
            else:
                return Response({
                    "status": True,
                    "message": "No Printers found",
                    "total": 0,
                    "data": []
                })
        except Exception as e:
            return Response({"status": False, 'message': 'Something went wrong', 'error': str(e)})


class PrinterMachinesUpdate(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    lookup_field = 'printer_id'

    def update(self, request, *args, **kwargs):
        try:
            printer_id = self.kwargs.get("printer_id")
            printer_name = request.data.get('printer_name')
            printer_description = request.data.get('printer_location')
    
            printer_obj = PrinterMachinesModel.objects.get(id=printer_id)
            if printer_name:
                printer_obj.printer_name = printer_name
            if printer_description:
                printer_obj.printer_description = printer_description
            printer_obj.save()
    
            return Response({'status': True, 'message': 'Print machine updated successfully'})
        except PrinterMachinesModel.DoesNotExist:
            return Response({'status': False, 'message': 'Print not found'})
        except Exception as e:
            return Response({'status': False, 'message': 'Something went wrong', 'error': str(e)})
        
    def destroy(self, request, *args, **kwargs):
        try:
            printer_id = request.data.get('printer_id')   

            if not PrinterMachinesModel.objects.get(id=printer_id):
                return Response({"status":False, "message":"Print machine id not found"})
                     
            printer_object = PrinterMachinesModel.objects.get(id=printer_id)
            printer_object.delete()
            return Response({"status":True, "message":"Printer deleted succesfully"})
        except Exception as e:
                return Response({"status": False,'message': 'Something went wrong','error': str(e)})


class DocumentReviseRequestViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = DocumentRevisionRequestAction.objects.all()

    def create(self, request, *args, **kwargs):
        try:
            user = self.request.user
            document_id = request.data.get('document_id')
            revise_description = request.data.get('revise_description')

            if not document_id:
                return Response({"status": False, "message": "Document ID is required"})
            if not revise_description:
                return Response({"status": False, "message": "Revise description is required"})

            document = Document.objects.get(id=document_id)

            revise_request = DocumentRevisionRequestAction.objects.filter(user=user, document=document).order_by('-created_at').first()

            if revise_request:
                revise_request.revise_description = revise_description
                revise_request.is_revise = True
                revise_request.save()
            else:
                revise_request = DocumentRevisionRequestAction.objects.create(
                    user=user,
                    document=document,
                    revise_description=revise_description,
                    is_revise=True
                )

            return Response({
                "status": True,
                "message": "Revise request created successfully",
                "data": {}
            })
        except Document.DoesNotExist:
            return Response({"status": False, "message": "Invalid document ID"})
        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})


class DocumentReviseRequestGetViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentSerializer

    def get_queryset(self):
        # Filter Document objects where current status ID is 7
        return Document.objects.filter(document_current_status_id=7)

    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset()
        serializer = self.get_serializer(queryset, many=True)

        # Get the group ID of the requested user
        user_groups = request.user.groups.values_list('id', flat=True)  # Get all group IDs
        group_id = user_groups[0] if user_groups else None  # Take the first group if available

        pending_revise_count = DocumentRevisionRequestAction.objects.filter(
            status="Pending"
        ).count()
        # Add is_revise field for each document
        for doc in serializer.data:
            is_revise = DocumentRevisionRequestAction.objects.filter(
                document_id=doc['document_id'], is_revise=True
            ).exists()
            doc['is_revise'] = is_revise  # Append is_revise to response

        return Response({
            "status": True,
            "message": "List of documents with current status ID 7 retrieved successfully",
            "user_group_id": group_id, 
            "pending_revise_count": pending_revise_count,
            "data": serializer.data
        })

class ApprovedPrintRequestViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = ApprovedPrintRequestSerializer

    def list(self, request, *args, **kwargs):
        """List only approved print requests (status ID = 9)."""
        try:
            queryset = PrintRequestApproval.objects.filter(status__id=9)

            if queryset.exists():
                serializer = self.get_serializer(queryset, many=True)
                return Response({
                    "status": True,
                    "message": "Approved print requests fetched successfully",
                    'total': queryset.count(),
                    'data': serializer.data
                })
            else:
                return Response({
                    "status": True,
                    "message": "No approved print requests found",
                    "total": 0,
                    "data": []
                })
        except Exception as e:
            return Response({"status": False, 'message': 'Something went wrong', 'error': str(e)})
        

class ApprovalNumberViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = ApprovalNumberSerializer
    queryset = ApprovalNumber.objects.all()

    def list(self, request, *args, **kwargs):
        print_request_approval_id = self.kwargs.get('print_request_approval_id')

        if not print_request_approval_id:
            return Response({"message": "PrintRequestApproval ID is required."})

        try:
            # Get the ApprovalNumbers related to the given PrintRequestApproval
            approval_numbers = ApprovalNumber.objects.filter(
                printrequestapproval__id=print_request_approval_id
            )

            # Get the retrival_numbers associated with the given PrintRequestApproval
            retrival_numbers = RetrivalNumber.objects.filter(
                printrequestapproval__id=print_request_approval_id
            ).values_list('retrival_number', flat=True)

            # Filter out the ApprovalNumbers whose number is already in RetrivalNumbers
            filtered_approval_numbers = approval_numbers.exclude(
                number__in=retrival_numbers
            )

            # Serialize the filtered ApprovalNumbers
            serializer = self.get_serializer(filtered_approval_numbers, many=True)

            return Response({
                'print_request_approval_id': print_request_approval_id,
                'approval_numbers': serializer.data
            })

        except PrintRequestApproval.DoesNotExist:
            return Response({"message": "PrintRequestApproval with the given ID does not exist."})


class PrintRequestApprovalViewSet(viewsets.ModelViewSet):
    queryset = PrintRequestApproval.objects.all()
    serializer_class = RetrivalNumberSerializer

    def create(self, request, *args, **kwargs):

        print_request_approval_id = request.data.get('print_request_approval_id')
        retrival_numbers_data = request.data.get('retrival_numbers', [])

        if not print_request_approval_id:
            return Response({"error": "print_request_approval_id is required"})
        if not isinstance(retrival_numbers_data, list):
            return Response({"error": "retrival_numbers should be a list"})

        try:
            print_request_approval = PrintRequestApproval.objects.get(id=print_request_approval_id)
        except PrintRequestApproval.DoesNotExist:
            return Response({"error": "PrintRequestApproval not found"})

        # Validate and create RetrivalNumber instances
        added_numbers = []
        for number in retrival_numbers_data:
            retrival_number, created = RetrivalNumber.objects.get_or_create(retrival_number=number)
            print_request_approval.retrival_numbers.add(retrival_number)
            added_numbers.append(retrival_number.retrival_number)

        return Response({
            "status": True,
            "message": "Retrival numbers added successfully",
        })

class RetrivalNumbersViewSet(viewsets.ModelViewSet):

    permission_classes = [permissions.IsAuthenticated]
    serializer_class = RetrivalNumberSerializer
    queryset = RetrivalNumber.objects.all()

    def list(self, request, *args, **kwargs):
        print_request_approval_id = self.kwargs.get('print_request_approval_id')

        if not print_request_approval_id:
            return Response({"message": "PrintRequestApproval ID is required."}, status=400)

        try:
            print_request_approval = PrintRequestApproval.objects.get(id=print_request_approval_id)

            retrival_numbers = print_request_approval.retrival_numbers.all()

            serializer = self.get_serializer(retrival_numbers, many=True)

            return Response({
                "print_request_approval_id": print_request_approval_id,
                "retrival_numbers": serializer.data
            })

        except PrintRequestApproval.DoesNotExist:
            return Response(
                {"message": "PrintRequestApproval with the given ID does not exist."}
            )
        
        
class DocumentwiseIdViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            document_id = self.kwargs.get('document_id', None)

            if document_id:
                queryset = Document.objects.filter(id=document_id)
            else:
                return Response({
                    "status": False,
                    "message": "Document ID is required",
                }, status=400)

            queryset = self.filter_queryset(queryset)

            if queryset.exists():
                serializer = self.serializer_class(queryset, many=True)
                return Response({
                    "status": True,
                    "message": "Documents fetched successfully",
                    'total': queryset.count(),
                    "data": serializer.data,
                })
            else:
                return Response({
                    "status": True,
                    "message": "No Documents found",
                    "total": 0,
                    "data": [],
                })
        except Exception as e:
            return Response({
                "status": False,
                'message': 'Something went wrong',
                'error': str(e)
            })

class AllDocumentViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = Document.objects.all()
    serializer_class = AllDocumentSerializer

    def list(self, request, *args, **kwargs):
        queryset = self.queryset.filter(document_type_id=1)
        serializer = self.get_serializer(queryset, many=True)
        return Response({
            "status": True,
            "message": "Documents fetched successfully",
            "total": queryset.count(),
            "data": serializer.data
        })

    
class ParentDocumentViewSet(viewsets.ModelViewSet):

    permission_classes = [permissions.IsAuthenticated]
    serializer_class = AllDocumentSerializer
    queryset = Document.objects.all()

    def list(self, request, *args, **kwargs):
        document_id = self.kwargs.get('document_id')
        
        # Filter queryset based on document_id
        if document_id:
            queryset = Document.objects.filter(parent_document_id=document_id).order_by('-id')
        else:
            queryset = Document.objects.none()  

        # Serialize the filtered queryset
        serializer = self.get_serializer(queryset, many=True)
        
        return Response({
            "status": True,
            "message": "Documents fetched successfully" if queryset.exists() else "No documents found",
            "total": queryset.count(),
            "data": serializer.data
        })

class DocumentTimelineViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    lookup_field = 'document_id'

    def list(self, request, *args, **kwargs):
        document_id = self.kwargs.get('document_id')

        if not document_id:
            return Response(
                {"error": "document_id is required in query parameters"})

        data = {}

        data['author_approvals'] = DocumentAuthorApproveActionSerializer(
            DocumentAuthorApproveAction.objects.filter(document_id=document_id), many=True
        ).data

        data['reviewer_actions'] = DocumentReviewerActionSerializer(
            DocumentReviewerAction.objects.filter(document_id=document_id), many=True
        ).data

        data['approver_actions'] = DocumentApproverActionSerializer(
            DocumentApproverAction.objects.filter(document_id=document_id), many=True
        ).data

        data['doc_admin_actions'] = DocumentDocAdminActionSerializer(
            DocumentDocAdminAction.objects.filter(document_id=document_id), many=True
        ).data

        data['send_back_actions'] = DocumentSendBackActionSerializer(
            DocumentSendBackAction.objects.filter(document_id=document_id), many=True
        ).data

        data['release_actions'] = DocumentReleaseActionSerializer(
            DocumentReleaseAction.objects.filter(document_id=document_id), many=True
        ).data

        data['effective_actions'] = DocumentEffectivenewActionSerializer(
            DocumentEffectiveAction.objects.filter(document_id=document_id), many=True
        ).data

        data['revision_actions'] = DocumentRevisionActionSerializer(
            DocumentRevisionAction.objects.filter(document_id=document_id), many=True
        ).data

        data['revision_requests'] = DocumentRevisionRequestActionSerializer(
            DocumentRevisionRequestAction.objects.filter(document_id=document_id), many=True
        ).data

        return Response(data)


class DocAdminUpdateViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = Document.objects.all()

    def create(self, request, *args, **kwargs):
        try:
            user = request.user
            document = request.data.get('document_id')
            status_id = request.data.get('status')

            if not document:
                return Response({"status": False, "message": "Document are required"})
            if not status_id:
                return Response({"status": False, "message": "Status is required"})
            
            document = Document.objects.get(id=document)
            status = DynamicStatus.objects.get(id=status_id)
            
            document_action = DocumentObsoleteAction.objects.create(
                user=user,
                document=document,
                status=status
            )

            document.document_current_status = status
            document.save() 
            return Response({"status": True, "message": "Status change to Obsolete successfully"})

        except DynamicStatus.DoesNotExist:
            return Response({"status": False, "message": "Invalid status ID"})
        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})




class DocumentCertificatePdfExportView(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    lookup_field = 'document_id'
    
    # Assuming you have a model called 'Document'
    # You can modify the query to your needs
    queryset = Document.objects.all()

    def list(self, request, *args, **kwargs):
        try:
            # Get the document based on the document_id (passed in URL kwargs)
            document_id = kwargs.get('document_id')
            document = self.queryset.get(id=document_id)
            
            # Create context dictionary for the template
            context = {
                'document_title': document.title,
                'document_number': document.number,
                'document_type': document.type,
                'version': document.version,
                'created_at': document.created_at,
                'document_current_status': document.current_status,
                'revision_date': document.revision_date,
                'effective_date': document.effective_date,
                'assigned_to_first_name': document.assigned_to.first_name,
                'assigned_to_last_name': document.assigned_to.last_name,
            }

            # Path to the HTML template for document cover page
            html_template_path = settings.BASE_DIR / 'templates' / 'document_cover_page.html'

            # Load the HTML template
            template = get_template(html_template_path)

            # Render the HTML with context data
            html = template.render(context)

            # Create a BytesIO buffer to write the PDF to memory
            buffer = BytesIO()

            # Use xhtml2pdf to generate the PDF from the HTML
            pisa_status = pisa.CreatePDF(html, dest=buffer)

            # Check if there were any errors during PDF generation
            if pisa_status.err:
                return Response({
                    'status': False,
                    'message': 'Error occurred while generating the PDF.',
                    'data': {}
                })

            # Get the timestamp for the PDF file name
            timestamp = int(time.time())
            file_name = f"document_{document_id}_{timestamp}.pdf"

            # Define the file path to save the PDF in 'document_certificare' folder
            file_path = os.path.join(settings.MEDIA_ROOT, 'document_certificare', file_name)

            # Create the 'document_certificare' folder if it doesn't exist
            folder_path = os.path.dirname(file_path)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)

            # Save the generated PDF file
            with open(file_path, 'wb') as f:
                f.write(buffer.getvalue())

            # Build the file URL
            pdf_file_url = f"{settings.MEDIA_URL}document_certificare/{file_name}"
            full_pdf_file_url = f"{request.scheme}://{request.get_host()}{pdf_file_url}"

            # Return the response with the PDF URL
            return Response({
                'status': True,
                'message': 'PDF file generated and saved successfully.',
                'data': {'file_url': full_pdf_file_url}
            })

        except Document.DoesNotExist:
            return Response({
                'status': False,
                'message': 'Document not found.',
                'data': {}
            })
        except Exception as e:
            return Response({
                'status': False,
                'message': str(e),
                'data': {}
            })
            


class GetDocumentCertificateDataListViewSet(viewsets.ViewSet):
    permission_classes = [permissions.IsAuthenticated]
    lookup_field = 'document_id'
    
    def list(self, request, *args, **kwargs):
        document_id = kwargs.get('document_id')
        try:
            # Fetch the document
            document = Document.objects.get(id=document_id)
            
            # Fetch all actions for the document
            all_actions = self.get_document_actions(document)
            
            # Prepare data to return in the response
            document_data = {
                'document_id': document.id,
                'document_title': document.document_title,
                'document_number': document.document_number,
                'version': document.version,
                'department': document.user.department.department_name,
                'effective_date': document.effective_date.strftime('%d-%m-%Y') if document.effective_date else None,
                'revision_date': document.revision_date.strftime('%d-%m-%Y') if document.revision_date else None,
            }
            
            # Collect actions data
            actions_data = []
            for action in all_actions:
                action_data = {
                    'role': action.user.groups.first().name if action.user.groups.exists() else "No Role",
                    'department': action.user.department.department_name if action.user.department else "No Department",
                    'designation': action.user.designation if action.user.designation else "No Designation",
                    'user_name': f"{action.user.first_name} {action.user.last_name}",
                    'created_at': action.created_at.strftime('%d-%m-%Y')
                }
                actions_data.append(action_data)

            # Combine document data and actions
            response_data = {
                'document': document_data,
                'actions': actions_data
            }
            return Response({"status": True,"message": "Data fetched successfully","data": response_data})

        except Document.DoesNotExist:
            return Response({"status": False,"message": "Document not found","data": ""})
        except Exception as e:
            return Response({"status": False,"message": str(e),"data": ""})

    def get_document_actions(self, document):
        actions = []
        action_models = [
            DocumentAuthorApproveAction,
            DocumentReviewerAction,
            DocumentApproverAction,
            DocumentDocAdminAction,
            DocumentSendBackAction,
            DocumentReleaseAction,
            DocumentEffectiveAction,
            DocumentRevisionAction,
        ]
        
        for model in action_models:
            actions.extend(model.objects.filter(document=document).order_by('created_at'))
        
        return actions

    def get_roles(self, user):
        roles = []
        if user.groups.exists():
            for group in user.groups.all():
                roles.append(group.name)
        else:
            roles.append("No Role")
        return roles


from PyPDF2 import PdfMerger
from docx2pdf import convert
import platform
class DocumentCertificatePdfExportView(viewsets.ViewSet):
    permission_classes = [permissions.IsAuthenticated]
    lookup_field = 'document_id'
    
    def list(self, request, *args, **kwargs):
        document_id = kwargs.get('document_id')
        try:
            # Fetch the document
            document = Document.objects.get(id=document_id)
            
            # Fetch all actions for the document
            all_actions = self.get_document_actions(document)

            # Fetch latest document comment for front_file_url
            latest_comment = NewDocumentCommentsData.objects.filter(document=document).order_by("-created_at").first()
            front_file_url = latest_comment.front_file_url.path if latest_comment and latest_comment.front_file_url else None
            print(front_file_url, "Front file")
            # Define the context for the template
            context = {
                'document': document,
                'all_actions': all_actions,
                'front_file_url': request.build_absolute_uri(latest_comment.front_file_url.url) if front_file_url else None,
                'logo': os.path.join(settings.BASE_DIR, 'static', 'certificate_logo_image', 'logo.png')
            }
            
            # Render the template with context data
            template = get_template('document_cover_page.html')
            html = template.render(context)

            timestamp = int(time.time())  # Timestamp in seconds
            base_filename = f"document_certificate_cover{timestamp}"
            pdf_filename = f"{base_filename}.pdf"
            pdf_path = os.path.join(settings.MEDIA_ROOT, 'document_cover', pdf_filename)
            os.makedirs(os.path.dirname(pdf_path), exist_ok=True)

            # Generate PDF from HTML
            with open(pdf_path, 'wb') as output_file:
                pisa_status = pisa.CreatePDF(html, dest=output_file)

            if pisa_status.err:
                return Response({"status": False, "message": "Error occurred while generating PDF", "data": ""})

            final_pdf_path = pdf_path  # Default to cover page PDF

            # Convert DOCX to PDF if the document exists
            if front_file_url and front_file_url.endswith('.docx'):
                docx_pdf_path = pdf_path.replace(".pdf", "_converted.pdf")
                self.convert_docx_to_pdf(front_file_url, docx_pdf_path)

                # Check if conversion was successful
                if os.path.exists(docx_pdf_path) and os.path.getsize(docx_pdf_path) > 0:
                    # Merge PDFs
                    merged_pdf_path = pdf_path.replace(".pdf", "_final.pdf")
                    merger = PdfMerger()
                    
                    merger.append(pdf_path)  # Main certificate PDF
                    merger.append(docx_pdf_path)  # Converted DOCX PDF
                    
                    merger.write(merged_pdf_path)
                    merger.close()

                    final_pdf_path = merged_pdf_path  # Use merged PDF as final file

            # Return the final merged PDF URL
            pdf_file_url = f"{settings.MEDIA_URL}document_cover/{os.path.basename(final_pdf_path)}"
            full_pdf_file_url = f"{request.scheme}://{request.get_host()}{pdf_file_url}"
            return Response({"status": True, "message": "PDF generated successfully", "data": full_pdf_file_url})

        except Document.DoesNotExist:
            return Response({"status": False, "message": "Document not found", "data": ""})
        except Exception as e:
            return Response({"status": False, "message": str(e), "data": ""})

    def convert_docx_to_pdf(self, docx_path, pdf_path):
        """Convert DOCX to PDF based on OS"""
        if platform.system() == "Windows":
            convert(docx_path, pdf_path)  # Use docx2pdf for Windows
        else:
            # Use LibreOffice for Linux
            try:
                subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(pdf_path), docx_path], check=True)
            except Exception as e:
                print(f"LibreOffice conversion failed: {e}")

    def get_document_actions(self, document):
        actions = []
        action_models = [
            DocumentAuthorApproveAction,
            DocumentReviewerAction,
            DocumentApproverAction,
            DocumentDocAdminAction,
            DocumentSendBackAction,
            DocumentReleaseAction,
            DocumentEffectiveAction,
            DocumentRevisionAction,
        ]

        # Get all reviewer actions (since multiple can exist)
        author_action = DocumentAuthorApproveAction.objects.filter(document=document).order_by('-created_at').first()
        if author_action:
            author_action.role = "Author"
            actions.append(author_action)
    
        # 2️⃣ Get all Reviewer actions (Multiple reviewers can exist)
        reviewer_actions = DocumentReviewerAction.objects.filter(document=document).order_by('created_at')
        for action in reviewer_actions:
            action.role = "Reviewer"
            actions.append(action)
    
        # 3️⃣ Get the Approver action (Only one latest approver)
        approver_action = DocumentApproverAction.objects.filter(document=document).order_by('-created_at').first()
        if approver_action:
            approver_action.role = "Approver"
            actions.append(approver_action)
    
        # 4️⃣ Get the Doc Admin action (Only one latest admin)
        doc_admin_action = DocumentDocAdminAction.objects.filter(document=document).order_by('-created_at').first()
        if doc_admin_action:
            doc_admin_action.role = "Doc Admin"
            actions.append(doc_admin_action)
    
        return actions




class DocumentNintyDaysDataViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            current_time = timezone.localtime(timezone.now())
            is_doc_admin = user.groups.filter(name="Doc Admin").exists()

            ninety_days_from_now = current_time + timedelta(days=90)
            documents = Document.objects.filter(revision_date__gte=current_time, revision_date__lte=ninety_days_from_now)

            if is_doc_admin:
                department = request.query_params.get('department_id', None)
                if department:
                    documents = documents.filter(user__department=department)
            else:
                user_department = user.department
                documents = documents.filter(user__department=user_department)

            document_count = documents.count()

            serializer = DocumentviewSerializer(documents, many=True, context={'request': request})
            data = serializer.data

            return Response({"status": True,"message": "Documents fetched successfully","data_count": document_count,"data": data})
        except Exception as e:
            return Response({"status": False,'message': 'Something went wrong','error': str(e)})
        
        
class ArchivedDocumentViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = Archived.objects.all()
    serializer_class = ArchivedSerializer

    def list(self, request, *args, **kwargs):
        document_id = self.kwargs.get('document_id')
        
        if not document_id:
            return Response({"error": "document_id is required"})
            
        try:
            document = Document.objects.get(id=document_id)
        except Document.DoesNotExist:
            return Response({'status': False, 'message': 'Document not found'})

        # Fetch the latest Archived entry for the given document
        archived_entry = Archived.objects.filter(document=document).latest('created_at')

        # Prepare the response data using the Archived serializer
        archived_data = ArchivedSerializer(archived_entry).data

        # Add additional fields from the Document model
        response_data = {
            "title": document.document_title,
            "type": document.document_type.document_name,  # Assuming document_type has a `name` field
            "document_number": document.document_number,
            "created_date": document.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            "status": document.document_current_status.status,  # Assuming status has `status_name` field
            "version": archived_data['version'],
        }

        return Response({'status': False,'message': 'Document not found.','data': {}})

class UpdateDocumentUserViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    queryset = Archived.objects.all()
    
    def update(self, request, *args, **kwargs):
        try:
            document_id = self.kwargs.get('document_id')
            visible_to_users = request.data.get('visible_to_users', [])
            approver = request.data.get('approver')
            doc_admin = request.data.get('doc_admin')

            # Ensure document_id is provided
            if not document_id:
                return Response({"status": False, "message": "Document ID is required"})

            # Fetch the document
            try:
                document = Document.objects.get(id=document_id)
            except Document.DoesNotExist:
                return Response({"status": False, "message": "Document not found"})

            # Update visible_to_users
            if isinstance(visible_to_users, str):
                import json
                try:
                    visible_to_users = json.loads(visible_to_users)
                except json.JSONDecodeError:
                    return Response({
                        "status": False,
                        "message": "Invalid format for visible_to_users. Provide a valid list of IDs.",
                    })
            if visible_to_users:
                document.visible_to_users.set(visible_to_users)

            # Update approver
            if approver:
                try:
                    approver_user = CustomUser.objects.get(id=approver)
                    document.approver = approver_user
                except CustomUser.DoesNotExist:
                    return Response({"status": False, "message": "Approver user not found"})

            # Update doc_admin
            if doc_admin:
                try:
                    doc_admin_user = CustomUser.objects.get(id=doc_admin)
                    document.doc_admin = doc_admin_user
                except CustomUser.DoesNotExist:
                    return Response({"status": False, "message": "Doc Admin user not found"})

            # Save updated document
            document.save()

            return Response({"status": True, "message": "Document updated successfully"})

        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            current_time = timezone.localtime(timezone.now())
            is_doc_admin = user.groups.filter(name="Doc Admin").exists()

            ninety_days_from_now = current_time + timedelta(days=90)
            documents = Document.objects.filter(revision_date__gte=current_time, revision_date__lte=ninety_days_from_now)

            if is_doc_admin:
                department = request.query_params.get('department_id', None)
                if department:
                    documents = documents.filter(user__department=department)
            else:
                user_department = user.department
                documents = documents.filter(user__department=user_department)

            document_count = documents.count()

            serializer = DocumentviewSerializer(documents, many=True, context={'request': request})
            data = serializer.data

            return Response({"status": True,"message": "Documents fetched successfully","data_count": document_count,"data": data})
        except Exception as e:
            return Response({"status": False,'message': 'Something went wrong','error': str(e)})
        
        

class DateWiseDocumentDatacountViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            is_doc_admin = user.groups.filter(name="Doc Admin").exists()
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)

            start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)

            if error:
                return Response({"status": False,"message": error,"data_count": 0,"data": []})

            if not start_date_obj or not end_date_obj:
                documents = Document.objects.all()
            else:
                documents = Document.objects.filter(created_at__gte=start_date_obj, created_at__lte=end_date_obj)

            if is_doc_admin:
                department = request.query_params.get('department_id', None)
                if department:
                    documents = documents.filter(user__department=department)
            else:
                user_department = user.department
                documents = documents.filter(user__department=user_department)

            document_count = documents.count()

            serializer = DocumentviewSerializer(documents, many=True, context={'request': request})
            data = serializer.data
            return Response({"status": True,"message": "Documents fetched successfully","data_count": document_count,"data": data})

        except Exception as e:
            return Response({"status": False,"message": "Something went wrong","error": str(e)})



class DocumentDataOfStatusIdOne(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            user_group_ids = user.groups.values_list('id', flat=True)
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)

            # Fixed status_id = 1 (You can change it to another ID if needed)
            fixed_status_id = 1

            # Fetch the status object for status_id = 1 (Fixed Status)
            status_obj = DynamicStatus.objects.filter(id=fixed_status_id).first()
            if not status_obj:
                return Response({"status": False, "message": "Status not found", "data": []})

            # Filter documents based on the fixed status_id
            queryset = Document.objects.filter(document_current_status=status_obj).order_by('-id')

            # Handle department-specific filtering for non-admin users
            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                if department_id:
                    queryset = queryset.filter(user__department_id=department_id)
            elif user.groups.filter(name='Reviewer').exists():
                queryset = queryset.filter(visible_to_users=user)
            elif user_department:
                queryset = queryset.filter(user__department=user_department)

            if start_date and end_date:
                start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
                if error:
                    return Response({"status": False, "message": error, "data": []})
    
                if start_date_obj:
                    start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
                if end_date_obj:
                    end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))
    
                if start_date_obj and end_date_obj:
                    queryset = queryset.filter(created_at__range=[start_date_obj, end_date_obj])

            document_count = queryset.count()

            serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
            if queryset.exists():
                return Response({"status": True,"message": "Documents fetched successfully","user_group_ids": list(user_group_ids),"data_count": document_count,"data": serializer.data})
            else:
                return Response({"status": True,"message": "No Documents found","user_group_ids": list(user_group_ids),"data": []})

        except Exception as e:
            return Response({"status": False,"message": "Something went wrong","error": str(e)})


class DocumentDataOfStatusIdTwo(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            user_group_ids = user.groups.values_list('id', flat=True)
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)

            fixed_status_id = 2
            status_obj = DynamicStatus.objects.filter(id=fixed_status_id).first()
            if not status_obj:
                return Response({"status": False, "message": "Status not found", "data": []})

            queryset = Document.objects.filter(document_current_status=status_obj).order_by('-id')

            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                if department_id:
                    queryset = queryset.filter(user__department_id=department_id)
            elif user.groups.filter(name='Reviewer').exists():
                queryset = queryset.filter(visible_to_users=user)
            elif user_department:
                queryset = queryset.filter(user__department=user_department)

            if start_date and end_date:
                start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
                if error:
                    return Response({"status": False, "message": error, "data": []})
    
                if start_date_obj:
                    start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
                if end_date_obj:
                    end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))
    
                if start_date_obj and end_date_obj:
                    queryset = queryset.filter(created_at__range=[start_date_obj, end_date_obj])

            document_count = queryset.count()

            serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
            if queryset.exists():
                return Response({"status": True,"message": "Documents fetched successfully","user_group_ids": list(user_group_ids),"data_count": document_count,"data": serializer.data})
            else:
                return Response({"status": True,"message": "No Documents found","user_group_ids": list(user_group_ids),"data": []})

        except Exception as e:
            return Response({"status": False,"message": "Something went wrong","error": str(e)})
        
class DocumentDataOfStatusIdThree(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            user_group_ids = user.groups.values_list('id', flat=True)
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)
    
            fixed_status_id = 3
            status_obj = DynamicStatus.objects.filter(id=fixed_status_id).first()
            if not status_obj:
                return Response({"status": False, "message": "Status not found", "data": []})
    
            # Apply filters BEFORE union
            queryset1 = Document.objects.filter(document_current_status=status_obj)
            queryset2 = Document.objects.filter(visible_to_users=user, document_current_status=8)
    
            # Apply department filter before union
            if department_id:
                queryset1 = queryset1.filter(user__department_id=department_id)
                queryset2 = queryset2.filter(user__department_id=department_id)
    
            # Apply date range filter before union
            if start_date and end_date:
                start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
                if error:
                    return Response({"status": False, "message": error, "data": []})
    
                if start_date_obj:
                    start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
                if end_date_obj:
                    end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))
    
                queryset1 = queryset1.filter(created_at__range=[start_date_obj, end_date_obj])
                queryset2 = queryset2.filter(created_at__range=[start_date_obj, end_date_obj])
    
            # Apply additional user-based filtering before union
            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                pass  # No extra filtering needed
            elif user.groups.filter(name='Reviewer').exists():
                queryset1 = queryset1.filter(visible_to_users=user)
                queryset2 = queryset2.filter(visible_to_users=user)
            elif user_department:
                queryset1 = queryset1.filter(user__department=user_department)
                queryset2 = queryset2.filter(user__department=user_department)
    
            # Perform union after all filtering
            queryset = queryset1.union(queryset2).order_by('-id')
    
            document_count = queryset.count()
    
            serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
            if queryset.exists():
                return Response({"status": True, "message": "Documents fetched successfully", "user_group_ids": list(user_group_ids), "data_count": document_count, "data": serializer.data})
            else:
                return Response({"status": True, "message": "No Documents found", "user_group_ids": list(user_group_ids), "data": []})
    
        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})

        
class DocumentDataOfStatusIdFour(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            user_group_ids = user.groups.values_list('id', flat=True)
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)

            fixed_status_id = 4
            status_obj = DynamicStatus.objects.filter(id=fixed_status_id).first()
            if not status_obj:
                return Response({"status": False, "message": "Status not found", "data": []})

            queryset = Document.objects.filter(document_current_status=status_obj).order_by('-id')

            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                if department_id:
                    queryset = queryset.filter(user__department_id=department_id)
            elif user.groups.filter(name='Reviewer').exists():
                queryset = queryset.filter(visible_to_users=user)
            elif user_department:
                queryset = queryset.filter(user__department=user_department)

            if start_date and end_date:
                start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
                if error:
                    return Response({"status": False, "message": error, "data": []})
    
                if start_date_obj:
                    start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
                if end_date_obj:
                    end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))
    
                if start_date_obj and end_date_obj:
                    queryset = queryset.filter(created_at__range=[start_date_obj, end_date_obj])

            document_count = queryset.count()

            serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
            if queryset.exists():
                return Response({"status": True,"message": "Documents fetched successfully","user_group_ids": list(user_group_ids),"data_count": document_count,"data": serializer.data})
            else:
                return Response({"status": True,"message": "No Documents found","user_group_ids": list(user_group_ids),"data": []})

        except Exception as e:
            return Response({"status": False,"message": "Something went wrong","error": str(e)})
        

class DocumentDataOfStatusIdFive(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            user_group_ids = user.groups.values_list('id', flat=True)
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)

            fixed_status_id = 5
            status_obj = DynamicStatus.objects.filter(id=fixed_status_id).first()
            if not status_obj:
                return Response({"status": False, "message": "Status not found", "data": []})

            queryset = Document.objects.filter(document_current_status=status_obj).order_by('-id')

            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                if department_id:
                    queryset = queryset.filter(user__department_id=department_id)
            elif user.groups.filter(name='Reviewer').exists():
                queryset = queryset.filter(visible_to_users=user)
            elif user_department:
                queryset = queryset.filter(user__department=user_department)

            if start_date and end_date:
                start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
                if error:
                    return Response({"status": False, "message": error, "data": []})
    
                if start_date_obj:
                    start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
                if end_date_obj:
                    end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))
    
                if start_date_obj and end_date_obj:
                    queryset = queryset.filter(created_at__range=[start_date_obj, end_date_obj])

            document_count = queryset.count()

            serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
            if queryset.exists():
                return Response({"status": True,"message": "Documents fetched successfully","user_group_ids": list(user_group_ids),"data_count": document_count,"data": serializer.data})
            else:
                return Response({"status": True,"message": "No Documents found","user_group_ids": list(user_group_ids),"data": []})

        except Exception as e:
            return Response({"status": False,"message": "Something went wrong","error": str(e)})
        

class DocumentDataOfStatusIdSix(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            user_group_ids = user.groups.values_list('id', flat=True)
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)

            fixed_status_id = 6
            status_obj = DynamicStatus.objects.filter(id=fixed_status_id).first()
            if not status_obj:
                return Response({"status": False, "message": "Status not found", "data": []})

            queryset = Document.objects.filter(document_current_status=status_obj).order_by('-id')

            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                if department_id:
                    queryset = queryset.filter(user__department_id=department_id)
            elif user.groups.filter(name='Reviewer').exists():
                queryset = queryset.filter(visible_to_users=user)
            elif user_department:
                queryset = queryset.filter(user__department=user_department)

            if start_date and end_date:
                start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
                if error:
                    return Response({"status": False, "message": error, "data": []})
    
                if start_date_obj:
                    start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
                if end_date_obj:
                    end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))
    
                if start_date_obj and end_date_obj:
                    queryset = queryset.filter(created_at__range=[start_date_obj, end_date_obj])

            document_count = queryset.count()

            serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
            if queryset.exists():
                return Response({"status": True,"message": "Documents fetched successfully","user_group_ids": list(user_group_ids),"data_count": document_count,"data": serializer.data})
            else:
                return Response({"status": True,"message": "No Documents found","user_group_ids": list(user_group_ids),"data": []})

        except Exception as e:
            return Response({"status": False,"message": "Something went wrong","error": str(e)})
        

class DocumentDataOfStatusIdSeven(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            user_group_ids = user.groups.values_list('id', flat=True)
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)

            fixed_status_id = 7
            status_obj = DynamicStatus.objects.filter(id=fixed_status_id).first()
            if not status_obj:
                return Response({"status": False, "message": "Status not found", "data": []})

            queryset = Document.objects.filter(document_current_status=status_obj).order_by('-id')

            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                if department_id:
                    queryset = queryset.filter(user__department_id=department_id)
            elif user.groups.filter(name='Reviewer').exists():
                queryset = queryset.filter(visible_to_users=user)
            elif user_department:
                queryset = queryset.filter(user__department=user_department)

            if start_date and end_date:
                start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
                if error:
                    return Response({"status": False, "message": error, "data": []})
    
                if start_date_obj:
                    start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
                if end_date_obj:
                    end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))
    
                if start_date_obj and end_date_obj:
                    queryset = queryset.filter(created_at__range=[start_date_obj, end_date_obj])

            document_count = queryset.count()

            serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
            if queryset.exists():
                return Response({"status": True,"message": "Documents fetched successfully","user_group_ids": list(user_group_ids),"data_count": document_count,"data": serializer.data})
            else:
                return Response({"status": True,"message": "No Documents found","user_group_ids": list(user_group_ids),"data": []})

        except Exception as e:
            return Response({"status": False,"message": "Something went wrong","error": str(e)})


class DocumentDataOfStatusIdEight(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            user_group_ids = user.groups.values_list('id', flat=True)
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)

            fixed_status_id = 8
            status_obj = DynamicStatus.objects.filter(id=fixed_status_id).first()
            if not status_obj:
                return Response({"status": False, "message": "Status not found", "data": []})

            queryset = Document.objects.filter(document_current_status=status_obj).order_by('-id')

            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                if department_id:
                    queryset = queryset.filter(user__department_id=department_id)
            elif user.groups.filter(name='Reviewer').exists():
                queryset = queryset.filter(visible_to_users=user)
            elif user_department:
                queryset = queryset.filter(user__department=user_department)

            if start_date and end_date:
                start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
                if error:
                    return Response({"status": False, "message": error, "data": []})
    
                if start_date_obj:
                    start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
                if end_date_obj:
                    end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))
    
                if start_date_obj and end_date_obj:
                    queryset = queryset.filter(created_at__range=[start_date_obj, end_date_obj])

            document_count = queryset.count()

            serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
            if queryset.exists():
                return Response({"status": True,"message": "Documents fetched successfully","user_group_ids": list(user_group_ids),"data_count": document_count,"data": serializer.data})
            else:
                return Response({"status": True,"message": "No Documents found","user_group_ids": list(user_group_ids),"data": []})

        except Exception as e:
            return Response({"status": False,"message": "Something went wrong","error": str(e)})
        
class DocumentDataOfStatusIdNine(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            user_group_ids = user.groups.values_list('id', flat=True)
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)

            fixed_status_id = 9
            status_obj = DynamicStatus.objects.filter(id=fixed_status_id).first()
            if not status_obj:
                return Response({"status": False, "message": "Status not found", "data": []})

            queryset = Document.objects.filter(document_current_status=status_obj).order_by('-id')

            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                if department_id:
                    queryset = queryset.filter(user__department_id=department_id)
            elif user.groups.filter(name='Reviewer').exists():
                queryset = queryset.filter(visible_to_users=user)
            elif user_department:
                queryset = queryset.filter(user__department=user_department)

            if start_date and end_date:
                start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
                if error:
                    return Response({"status": False, "message": error, "data": []})
    
                if start_date_obj:
                    start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
                if end_date_obj:
                    end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))
    
                if start_date_obj and end_date_obj:
                    queryset = queryset.filter(created_at__range=[start_date_obj, end_date_obj])

            document_count = queryset.count()

            serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
            if queryset.exists():
                return Response({"status": True,"message": "Documents fetched successfully","user_group_ids": list(user_group_ids),"data_count": document_count,"data": serializer.data})
            else:
                return Response({"status": True,"message": "No Documents found","user_group_ids": list(user_group_ids),"data": []})

        except Exception as e:
            return Response({"status": False,"message": "Something went wrong","error": str(e)})
        

class DocumentDataOfStatusIdTen(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            user_group_ids = user.groups.values_list('id', flat=True)
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)

            fixed_status_id = 10
            status_obj = DynamicStatus.objects.filter(id=fixed_status_id).first()
            if not status_obj:
                return Response({"status": False, "message": "Status not found", "data": []})

            queryset = Document.objects.filter(document_current_status=status_obj).order_by('-id')

            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                if department_id:
                    queryset = queryset.filter(user__department_id=department_id)
            elif user.groups.filter(name='Reviewer').exists():
                queryset = queryset.filter(visible_to_users=user)
            elif user_department:
                queryset = queryset.filter(user__department=user_department)

            if start_date and end_date:
                start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
                if error:
                    return Response({"status": False, "message": error, "data": []})
    
                if start_date_obj:
                    start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
                if end_date_obj:
                    end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))
    
                if start_date_obj and end_date_obj:
                    queryset = queryset.filter(created_at__range=[start_date_obj, end_date_obj])

            document_count = queryset.count()

            serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
            if queryset.exists():
                return Response({"status": True,"message": "Documents fetched successfully","user_group_ids": list(user_group_ids),"data_count": document_count,"data": serializer.data})
            else:
                return Response({"status": True,"message": "No Documents found","user_group_ids": list(user_group_ids),"data": []})

        except Exception as e:
            return Response({"status": False,"message": "Something went wrong","error": str(e)})
        


class DocumentDataOfStatusIdEleven(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            user_group_ids = user.groups.values_list('id', flat=True)
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)

            fixed_status_id = 11
            status_obj = DynamicStatus.objects.filter(id=fixed_status_id).first()
            if not status_obj:
                return Response({"status": False, "message": "Status not found", "data": []})

            queryset = Document.objects.filter(document_current_status=status_obj).order_by('-id')

            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                if department_id:
                    queryset = queryset.filter(user__department_id=department_id)
            elif user.groups.filter(name='Reviewer').exists():
                queryset = queryset.filter(visible_to_users=user)
            elif user_department:
                queryset = queryset.filter(user__department=user_department)

            if start_date and end_date:
                start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
                if error:
                    return Response({"status": False, "message": error, "data": []})
    
                if start_date_obj:
                    start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
                if end_date_obj:
                    end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))
    
                if start_date_obj and end_date_obj:
                    queryset = queryset.filter(created_at__range=[start_date_obj, end_date_obj])

            document_count = queryset.count()

            serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
            if queryset.exists():
                return Response({"status": True,"message": "Documents fetched successfully","user_group_ids": list(user_group_ids),"data_count": document_count,"data": serializer.data})
            else:
                return Response({"status": True,"message": "No Documents found","user_group_ids": list(user_group_ids),"data": []})

        except Exception as e:
            return Response({"status": False,"message": "Something went wrong","error": str(e)})
        

class DocumentDataOfStatusIdTwelve(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            user_group_ids = user.groups.values_list('id', flat=True)
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)

            fixed_status_id = 12
            status_obj = DynamicStatus.objects.filter(id=fixed_status_id).first()
            if not status_obj:
                return Response({"status": False, "message": "Status not found", "data": []})

            queryset = Document.objects.filter(document_current_status=status_obj).order_by('-id')

            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                if department_id:
                    queryset = queryset.filter(user__department_id=department_id)
            elif user.groups.filter(name='Reviewer').exists():
                queryset = queryset.filter(visible_to_users=user)
            elif user_department:
                queryset = queryset.filter(user__department=user_department)

            if start_date and end_date:
                start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
                if error:
                    return Response({"status": False, "message": error, "data": []})
    
                if start_date_obj:
                    start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
                if end_date_obj:
                    end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))
    
                if start_date_obj and end_date_obj:
                    queryset = queryset.filter(created_at__range=[start_date_obj, end_date_obj])

            document_count = queryset.count()

            serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
            if queryset.exists():
                return Response({"status": True,"message": "Documents fetched successfully","user_group_ids": list(user_group_ids),"data_count": document_count,"data": serializer.data})
            else:
                return Response({"status": True,"message": "No Documents found","user_group_ids": list(user_group_ids),"data": []})

        except Exception as e:
            return Response({"status": False,"message": "Something went wrong","error": str(e)})
        

class DocumentDataOfStatusIdThirteen(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all().order_by('-id')

    def list(self, request, *args, **kwargs):
        try:
            user = self.request.user
            user_group_ids = user.groups.values_list('id', flat=True)
            user_department = user.department if user.department else None
            department_id = request.query_params.get('department_id', None)
            start_date = request.query_params.get('start_date', None)
            end_date = request.query_params.get('end_date', None)

            fixed_status_id = 13
            status_obj = DynamicStatus.objects.filter(id=fixed_status_id).first()
            if not status_obj:
                return Response({"status": False, "message": "Status not found", "data": []})

            queryset = Document.objects.filter(document_current_status=status_obj).order_by('-id')

            if user.groups.filter(name='Admin').exists() or user.groups.filter(name='Doc Admin').exists():
                if department_id:
                    queryset = queryset.filter(user__department_id=department_id)
            elif user.groups.filter(name='Reviewer').exists():
                queryset = queryset.filter(visible_to_users=user)
            elif user_department:
                queryset = queryset.filter(user__department=user_department)

            if start_date and end_date:
                start_date_obj, end_date_obj, error = validate_dates(start_date, end_date)
                if error:
                    return Response({"status": False, "message": error, "data": []})
    
                if start_date_obj:
                    start_date_obj = timezone.make_aware(datetime.combine(start_date_obj, datetime.min.time()))
                if end_date_obj:
                    end_date_obj = timezone.make_aware(datetime.combine(end_date_obj, datetime.max.time()))
    
                if start_date_obj and end_date_obj:
                    queryset = queryset.filter(created_at__range=[start_date_obj, end_date_obj])

            document_count = queryset.count()

            serializer = DocumentviewSerializer(queryset, many=True, context={'request': request})
            if queryset.exists():
                return Response({"status": True,"message": "Documents fetched successfully","user_group_ids": list(user_group_ids),"data_count": document_count,"data": serializer.data})
            else:
                return Response({"status": True,"message": "No Documents found","user_group_ids": list(user_group_ids),"data": []})

        except Exception as e:
            return Response({"status": False,"message": "Something went wrong","error": str(e)})
import json
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import jwt
import os
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import hashlib
from django.http import FileResponse
from django.shortcuts import redirect

@csrf_exempt
# def get_editor_config(request):
#     document_url = "http://host.docker.internal:8000/media/templates/SOP_template.docx"  # Ensure this URL is valid and accessible 
#     unique_key = hashlib.sha256(document_url.encode()).hexdigest()

#     document_data = {
#         "fileType": "docx",
#         "key": unique_key,
#         "title": "Sample Document",
#         "url": document_url,  # Direct link to the document
#     }

#     editor_config = {
        
#         "document": document_data,
#         "editorConfig": {
#             # "callbackUrl": "http://127.0.0.1:8000/dms_module/onlyoffice_callback",
#             "callbackUrl": "http://host.docker.internal:8000/dms_module/onlyoffice_callback",
#             "mode": "edit",
#             "user": {"id": "1", "name": "Rohit Sharma"},
#         },
#     }

#     # Secret must match ONLYOFFICE configuration
#     secret = "45540a6bfecc97ab6d06c436a74c333b1b54447c4de5fd41b8ad0b8361a395c6"
#     token = jwt.encode(editor_config, secret, algorithm="HS256")

#     return JsonResponse({"token": token, **editor_config})
@csrf_exempt
def get_editor_config_for_obsolete_doc(request):
    # Get template_id from the request parameters
    # document_id = request.GET.get('document_id')
    # template_id = request.GET.get('template_id')
    front_file_url = request.GET.get('front_file_url') 

    # if not document_id:
    #     return JsonResponse({"status": False, "message": "document_id parameter is required"})
    # if not template_id:
    #     return JsonResponse({"status": False, "message": "template_id parameter is required"})
    if not front_file_url:
        return JsonResponse({"status": False, "message": "front_file_url parameter is required"})

    try:
        # doc = Document.objects.filter(id=document_id).first()
        
        # Fetch the latest document associated with the template_id
        # document = Document.objects.filter(select_template_id=template_id).order_by('-created_at').first()
        # document_name = Document.objects.filter(id=document_id).first()
        # print(document,"===========")
        # if not document:
        #     return JsonResponse({"status": False, "message": "No document found for the selected template"})

        # Get the document file URL
        # if not document.generatefile:
        #     return JsonResponse({"status": False, "message": "Document file not available for the selected template"})
        # BASE_URL = "http://host.docker.internal:8000"
        # BASE_URL ="http://13.232.63.196:8080"
        # Construct full document URL (assuming media files are served under MEDIA_URL)
        # document_url = f"{BASE_URL}{settings.MEDIA_URL}/generated_docs/{document.generatefile}"  # Ensure MEDIA_URL is properly configured
        # latest_comment = NewDocumentCommentsData.objects.filter(document=doc).order_by("-created_at").first()
        # front_file_url_ = latest_comment.front_file_url.url if latest_comment and latest_comment.front_file_url else None
        # front_file_url = f"{BASE_URL}{front_file_url_}"
        # if front_file_url == 'http://host.docker.internal:8000None':
        # if front_file_url == 'http://13.232.63.196:8080None':
        #     unique_key = hashlib.sha256(document_url.encode()).hexdigest()
        # else:
        unique_key = hashlib.sha256(front_file_url.encode()).hexdigest()
        # print( "document_url",document_url)
        # Document data
        document_data = {
            "fileType": "docx",  # Assuming the document is of type 'docx'
            "key": unique_key,
            "title": "Document",  # Use the document title
            # "url": document_url if front_file_url == 'http://host.docker.internal:8000None' else front_file_url,  # Direct link to the document
            # "url": document_url if front_file_url == 'http://13.232.63.196:8080None' else front_file_url,
            "url": front_file_url
        }

        # Editor configuration data
        editor_config = {
            "document": document_data,
            "editorConfig": {
                # Replace with your actual callback URL for the editor
                # "callbackUrl": "http://host.docker.internal:8080/dms_module/onlyoffice_callback",
                "callbackUrl": "http://13.232.63.196:8080/dms_module/onlyoffice_callback",
                "mode": "edit",
                "user": {"id": "1", "name": "Rohit Sharma"},
            },
        }

        # Secret for encoding the JWT token
        secret = "45540a6bfecc97ab6d06c436a74c333b1b54447c4de5fd41b8ad0b8361a395c6"
        token = jwt.encode(editor_config, secret, algorithm="HS256")

        return JsonResponse({"token": token, **editor_config})

    except Exception as e:
        return JsonResponse({"status": False, "message": str(e)})
    
@csrf_exempt
def get_editor_config(request):
    # Get template_id from the request parameters
    document_id = request.GET.get('document_id')
    template_id = request.GET.get('template_id')
    is_view = request.GET.get('is_view')

    if not document_id:
        return JsonResponse({"status": False, "message": "document_id parameter is required"})
    if not template_id:
        return JsonResponse({"status": False, "message": "template_id parameter is required"})

    try:
        doc = Document.objects.filter(id=document_id).first()
        user_agent = request.META.get('HTTP_USER_AGENT', '').lower()
        if "print" in user_agent or request.GET.get('print_mode') == "true":
            return JsonResponse({"status": False, "message": "Printing is restricted.", "document_content": ""})
        
        # Fetch the latest document associated with the template_id
        document = Document.objects.filter(select_template_id=template_id).order_by('-created_at').first()
        document_name = Document.objects.filter(id=document_id).first()
        print(document,"===========")
        if not document:
            return JsonResponse({"status": False, "message": "No document found for the selected template"})

        # Get the document file URL
        if not document.generatefile:
            return JsonResponse({"status": False, "message": "Document file not available for the selected template"})
        # BASE_URL = "http://host.docker.internal:8000"
        BASE_URL ="http://13.232.63.196:8080"
        # Construct full document URL (assuming media files are served under MEDIA_URL)
        document_url = f"{BASE_URL}{settings.MEDIA_URL}/generated_docs/{document.generatefile}"  # Ensure MEDIA_URL is properly configured
        latest_comment = NewDocumentCommentsData.objects.filter(document=doc).order_by("-created_at").first()
        front_file_url_ = latest_comment.front_file_url.url if latest_comment and latest_comment.front_file_url else None
        front_file_url = f"{BASE_URL}{front_file_url_}"
        # if front_file_url == 'http://host.docker.internal:8000None':
        if front_file_url == 'http://13.232.63.196:8080None':
            unique_key = hashlib.sha256(document_url.encode()).hexdigest()
        else:
            unique_key = hashlib.sha256(front_file_url.encode()).hexdigest()
        print( "document_url",document_url)
        is_view = str(is_view).lower() in ["true", "1"]
        # Document data
        document_data = {
            "fileType": "docx",  # Assuming the document is of type 'docx'
            "key": unique_key,
            "title": document_name.document_title or "Untitled Document",  # Use the document title
            # "url": document_url if front_file_url == 'http://host.docker.internal:8000None' else front_file_url,  # Direct link to the document
            "url": document_url if front_file_url == 'http://13.232.63.196:8080None' else front_file_url,
            "permissions": {
                "print": False,
                "download": not is_view
            }
        }

        # Editor configuration data
        editor_config = {
            "document": document_data,
            "editorConfig": {
                # Replace with your actual callback URL for the editor
                # "callbackUrl": "http://host.docker.internal:8080/dms_module/onlyoffice_callback",
                "callbackUrl": "http://13.232.63.196:8080/dms_module/onlyoffice_callback",
                "mode": "edit",
                "user": {"id": "1", "name": "Rohit Sharma"},
            },
        }

        # Secret for encoding the JWT token
        secret = "45540a6bfecc97ab6d06c436a74c333b1b54447c4de5fd41b8ad0b8361a395c6"
        token = jwt.encode(editor_config, secret, algorithm="HS256")

        return JsonResponse({"token": token, **editor_config})

    except Exception as e:
        return JsonResponse({"status": False, "message": str(e)})
@csrf_exempt
def onlyoffice_callback(request):
    print(f"Received request: {request.method} {request.body}")
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            print("Parsed data:", data)
            return JsonResponse({"error": 0})
        except Exception as e:
            print("Error:", e)
            return JsonResponse({"error": 1}, status=500)
    return JsonResponse({"error": 1}, status=400)



class EmployeeJobRoleView(viewsets.ViewSet):
    permission_classes = [permissions.IsAuthenticated]
    lookup_field = 'employee_id'

    def list(self, request, *args, **kwargs):
        employee_id = kwargs.get('employee_id')
        try:
            employee = CustomUser.objects.get(id=employee_id)
            department = Department.objects.filter(id=employee.department_id).first()
            job_assign = JobAssign.objects.filter(user=employee).first()
            job_roles = ', '.join(job_assign.job_roles.values_list('job_role_name', flat=True)) if job_assign else 'N/A'
            context = {
                'employee_number': employee.employee_number,
                'employee_name': f"{employee.first_name} {employee.last_name}",
                'designation': employee.designation,
                'department': department.department_name if employee.department else 'N/A',
                'joining_date': employee.created_at.strftime('%d/%m/%Y'),
                'company_name': 'Promount Technologies LLP',
                'department_name': department.department_name if employee.department else 'N/A',
                'job_role': job_roles
            }
            template = get_template('employee_job_role.html')
            html = template.render(context)
            timestamp = int(time.time())  # Timestamp in seconds
            filename = f"employee_job_role{timestamp}.pdf"
            file_path = os.path.join(settings.MEDIA_ROOT, 'employee_report', filename)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            with open(file_path, 'wb') as output_file:
                pisa_status = pisa.CreatePDF(html, dest=output_file)
            if pisa_status.err:
                return Response({"status": False, "message": "Error occurred while generating PDF", "data": ""})
            
            pdf_file_url = f"{settings.MEDIA_URL}employee_report/{filename}"
            full_pdf_file_url = f"{request.scheme}://{request.get_host()}{pdf_file_url}"
            return Response({"status": True, "message": "PDF generated successfully", "data": full_pdf_file_url})

        except Document.DoesNotExist:
            return Response({"status": False, "message": "Document not found", "data": ""})
        except Exception as e:
            return Response({"status": False, "message": str(e), "data": ""})


class EmployeeRecordLogView(viewsets.ViewSet):
    permission_classes = [permissions.IsAuthenticated]
    
    def list(self, request, *args, **kwargs):
        try:
            users = CustomUser.objects.all()

            all_users_data = []
            for user in users:
                if user.department:
                    department_name = user.department.department_name
                else:
                    department_name = "No Department"
                datestatus = QuizSession.objects.filter(user=user).first()
                name = TrainingCreate.objects.filter(created_by=user).first()
                document_number = Document.objects.filter(user=user).first()
                version = document_number.version if document_number else "No Version"
                trainer = Trainer.objects.filter(user=user).first()

                training_date = datestatus.started_at if datestatus else "Not started"
                status = datestatus.status if datestatus else "No Status"
                training_name = name.training_name if name else "No Training"
                document_number = document_number.document_number if document_number else "No Document"
                trainer_name = trainer.trainer_name if trainer else "No Trainer"

                user_data  = {
                    'employee_name': user.username,
                    'designation': user.designation,
                    'department_name': department_name,
                    'training_date': training_date,
                    'training_name': training_name,
                    'status': status,
                    'document_number': document_number,
                    'current_version': version,
                    'trainer_name': trainer_name,
                }
                print(user_data)
                all_users_data.append(user_data)
            context = {'users_data': all_users_data}

            template = get_template('employee_record_log.html')
            html = template.render(context)
            print(html)
            timestamp = int(time.time())
            filename = f"employee_record_log{timestamp}.pdf"
            file_path = os.path.join(settings.MEDIA_ROOT, 'employee_record_log', filename)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            with open(file_path, 'wb') as output_file:
                pisa_status = pisa.CreatePDF(html, dest=output_file)
            if pisa_status.err:
                return Response({"status": False, "message": "Error occurred while generating PDF", "data": ""})
            
            pdf_file_url = f"{settings.MEDIA_URL}employee_record_log/{filename}"
            full_pdf_file_url = f"{request.scheme}://{request.get_host()}{pdf_file_url}"
            return Response({"status": True, "message": "PDF generated successfully", "data": full_pdf_file_url})

        except Document.DoesNotExist:
            return Response({"status": False, "message": "Document not found", "data": ""})
        except Exception as e:
            return Response({"status": False, "message": str(e), "data": ""})


class EmployeeTrainingNeedIdentyView(viewsets.ViewSet):
    permission_classes = [permissions.IsAuthenticated]
    lookup_field = 'employee_id'

    def list(self, request, *args, **kwargs):
        employee_id = kwargs.get('employee_id')
        try:
            employee = CustomUser.objects.get(id=employee_id)
            user_job_roles = JobRole.objects.filter(job_assigns__user=employee)
            print(user_job_roles, "fffffffffffff")
            trainings = Document.objects.filter(job_roles__in=user_job_roles).distinct()
            print(trainings, "fffffffffffff")
            failed_quiz_sessions = AttemptedQuiz.objects.filter(user=employee, is_pass=False)
            # if not failed_quiz_sessions.exists():
            #     return Response({"status": True, "message": "User has passed all quizzes"})
            print(failed_quiz_sessions,"gggggggggggggg")
            users = CustomUser.objects.filter(id__in=failed_quiz_sessions.values_list('user_id', flat=True))
            print(users,"ffffffvvvvv")
            all_users_data = []
            for user in users:
                if user.department:
                    department_name = user.department.department_name
                else:
                    department_name = "No Department"

                user_failed_sessions = failed_quiz_sessions.filter(user=user)
                user_training_data = []
                
                for training in trainings:
                    if failed_quiz_sessions.filter(document=training).exists():
                        user_training_data.append({
                            'training_name': training.document_title,
                            'training_number': training.document_number
                        })
                print(user_training_data)
                status = "Pass" if user_failed_sessions.first().is_pass else "Failed"
                
                user_data = {
                    'employee_name': user.username,
                    'designation': user.designation,
                    'department_name': department_name,
                    'status': status,
                    'trainings': user_training_data, 
                }
                all_users_data.append(user_data)

            context = {'users_data': all_users_data}
            template = get_template('training_need_identy.html')
            html = template.render(context)

            timestamp = int(time.time())
            filename = f"employee_training_need_identy{timestamp}.pdf"
            file_path = os.path.join(settings.MEDIA_ROOT, 'employee_training_need_identy', filename)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            with open(file_path, 'wb') as output_file:
                pisa_status = pisa.CreatePDF(html, dest=output_file)

            if pisa_status.err:
                return Response({"status": False, "message": "Error occurred while generating PDF", "data": ""})

            pdf_file_url = f"{settings.MEDIA_URL}employee_training_need_identy/{filename}"
            full_pdf_file_url = f"{request.scheme}://{request.get_host()}{pdf_file_url}"
            return Response({"status": True, "message": "PDF generated successfully", "data": full_pdf_file_url})

        except CustomUser.DoesNotExist:
            return Response({"status": False, "message": "User not found", "data": ""})
        except Exception as e:
            return Response({"status": False, "message": str(e), "data": ""})
        
@csrf_exempt
def download_and_save_docx(request):
    url = request.POST.get('url')  

    if not url:
        return JsonResponse({"status": False, "message": "URL parameter is required"})

    try:
        response = requests.get(url)

        document = DocumentLink()
        document.docxfile.save(f"{uuid4()}.docx", ContentFile(response.content))

        document.save()
        local_url = request.build_absolute_uri(settings.MEDIA_URL)
        file_url = f"{local_url}{document.docxfile.name}"

        return JsonResponse({"status": True, "file_url": file_url})

    except requests.RequestException as e:
        return JsonResponse({"status": False, "message": f"Error downloading file: {str(e)}"})
    except Exception as e:
        return JsonResponse({"status": False, "message": f"An error occurred: {str(e)}"})
    
class DocumentObsoleteNotificationViewSet(viewsets.ViewSet):
    permission_classes = [permissions.IsAuthenticated]
    def list(self, request, *args, **kwargs):
        try:
            obsolete_documents = Document.objects.filter(document_current_status__id=3)

            if not obsolete_documents.exists():
                return Response({"status": False,"message": "No obsolete documents found."})

            users_notified = []

            for document in obsolete_documents:
                job_roles = document.job_roles.all()
                if not job_roles.exists():
                    return Response({"status": False,"message": "Job roles not found."})
                
                job_role_ids = job_roles.values_list('id', flat=True)
                users = JobAssign.objects.filter(job_roles__in=job_roles).values_list('user_id', flat=True).distinct()
                
                if not users.exists():
                    continue 

                users = CustomUser.objects.filter(id__in=users)
                for user in users:
                    print(f"User: {user.username}, Email: {user.email}")
                user_emails = set(user.email for user in users if user.email)
                
                if user_emails:
                    subject = f"Document {document.document_title} is Obsolete"
                    message = (
                        f"The document '{document.document_title}' with document number "
                        f"'{document.document_number}' has been marked as obsolete. "
                        "Please review the document and take necessary actions."
                    )
                    from_email = settings.DEFAULT_FROM_EMAIL

                    send_mail(
                        subject, message, from_email, list(user_emails), fail_silently=False
                    )

                    users_notified.append({
                        "document_id": document.id,
                        "document_title": document.document_title,
                        "emails_notified": list(user_emails)
                    })

            if not users_notified:
                return Response({"status": False,"message": "No users to notify for the obsolete documents.",})

            return Response({
                "status": True,
                "message": "Emails sent successfully to users assigned to obsolete documents.",
                "data": users_notified
            })

        except Exception as e:
            return Response({"status": False, "message": "Something went wrong", "error": str(e)})
        



class AddNewDocumentCommentsdataViewSet(viewsets.ModelViewSet):
    queryset = NewDocumentCommentsData.objects.all().order_by('-id')
    serializer_class = AddNewDocumentCommentsSerializer
    permission_classes = [permissions.IsAuthenticated]

    def create(self, request, *args, **kwargs):
        try:
            user = self.request.user
            document_id = request.data.get('document_id')
            comment_data = request.data.get('comment_data')
            version_no = request.data.get('version_no')
            front_file_url = request.data.get('front_file_url')
            # template_id = request.data.get('template_id')
            department_id = request.data.get('department_id')

            # Check if front_file_url is provided
            if not front_file_url:
                return Response({"status": False, "message": "front_file_url parameter is required", "data": []})

            # Download the file from the provided URL
            response = requests.get(front_file_url)

            if response.status_code != 200:
                return Response({"status": False, "message": "Failed to download the file from the provided URL", "data": []})

            # Save the document file locally
            document = NewDocumentCommentsData()
            # Save the file with a unique name using uuid4
            document.front_file_url.save(f"{uuid4()}.docx", ContentFile(response.content))
            try:
                document_data = Document.objects.filter(id=document_id).first()
            except Document.DoesNotExist:
                return Response({"status": False, "message": "Document not found", "data": []})
            
            # try:
            #     template_data = TemplateModel.objects.get(id=template_id)
            # except TemplateModel.DoesNotExist:
            #     return Response({"status": False, "message": "Template not found", "data": []})
            
            try:
                department_data = Department.objects.get(id=department_id)
            except Department.DoesNotExist:
                return Response({"status": False, "message": "Department not found", "data": []})


            # Create the NewDocumentCommentsData object
            created = NewDocumentCommentsData.objects.create(
                user=user,
                document=document_data,
                comment_data=comment_data,
                version_no=version_no,
                front_file_url=document.front_file_url.name,
                department_id=department_data.id
            )

            # Serialize the created data
            serializer = AddNewDocumentCommentsSerializer(created)

            # Generate the absolute URL for the saved file
            local_url = request.build_absolute_uri(settings.MEDIA_URL)
            file_url = f"{local_url}{document.front_file_url.name}"

            return Response({
                "status": True,
                "message": "Document comments created successfully",
                "data": serializer.data,
                "file_url": file_url  # Return the new URL for the saved file
            })

        except Exception as e:
            return Response({"status": False, "message": str(e), "data": []})


class DocumentEffectiveViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentEffectiveSerializer

    def create(self, request, *args, **kwargs):
        try:
            user = request.user
            document_id = request.data.get('document')
            status = request.data.get('status')

            if not document_id or not status:
                return Response({"status": False, "message": "document_id and effective_date are required", "data": []})
        
            try:
                document_data = Document.objects.filter(id=document_id).first()
            except Document.DoesNotExist:
                return Response({"status": False, "message": "Document not found", "data": []})
            try:
                status_data = DynamicStatus.objects.get(id=status)
            except DynamicStatus.DoesNotExist:
                return Response({"status": False, "message": "Status not found", "data": []})
            
            document_data.document_current_status = status_data
            document_data.effective_date = datetime.now()
            document_data.save()
            if document_data.version and document_data.version.endswith(".0"): 
                previous_major_version = str(int(document_data.version.split(".")[0]) - 1) + ".0"

                old_doc = Document.objects.filter(
                    document_number=document_data.document_number,
                    version=previous_major_version
                ).update(document_current_status=DynamicStatus.objects.get(id=15))
                
            return Response({"status": True,"message": "Document effective date with supersede created successfully"})
        
        except Exception as e:
            return Response({"status": False, "message": str(e), "data": []})


class DocumentVersionListViewSet(viewsets.ModelViewSet):
    serializer_class = DocumentSerializer
    permission_classes = [permissions.IsAuthenticated]


    def list(self, request, *args, **kwargs):
        
        queryset = Document.objects.filter(document_current_status__in = [15,12])
        
        if not queryset.exists():
            return Response({"status": True, "message": "No comments found for this document", "data": []})

        serializer = self.serializer_class(queryset, many=True, context={'request': request})

        return Response({
            "status": True,
            "message": "Document comments fetched successfully",
            "data": serializer.data
        })


class DocxConvertPDFViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    
    def update(self, request, *args, **kwargs):
        try:
            document_id = self.kwargs.get('document_id')
            document_instance = Document.objects.get(id=document_id)
            latest_comment = NewDocumentCommentsData.objects.filter(document=document_instance).order_by('-created_at').first()
            url = latest_comment.front_file_url
            
            if not url.endswith('.docx'):
                return Response({'status': False, 'message': 'Invalid document type. Only .docx files are supported.'})

            base_directory = os.path.join(settings.MEDIA_ROOT, 'generated_docs')
            docx_file_path = os.path.join(base_directory, url)

            if not os.path.exists(docx_file_path):
                return Response({'status': False, 'message': f"Document file not found at {docx_file_path}."})

            try:
                pdf_output_path = docx_file_path.replace('.docx', '.pdf')
                convert(docx_file_path)

                pdf_relative_path = os.path.relpath(pdf_output_path, settings.MEDIA_ROOT)
                pdf_url = f"{settings.MEDIA_URL}{pdf_relative_path}"

            except Exception as e:
                return Response({'status': False, 'message': 'Error during conversion.', 'error': str(e)})

            return Response({
                'status': True,
                'message': 'Document successfully converted to PDF.',
                'pdf_link': request.build_absolute_uri(pdf_url)
            })

        except PrintRequest.DoesNotExist:
            return Response({'status': False, 'message': 'Print request not found.'})
        except Exception as e:
            return Response({'status': False, 'message': 'An error occurred while processing the document.', 'error': str(e)})


class IDWiSeDocumentListViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DocumentviewSerializer
    queryset = Document.objects.all()

    def list(self, request, *args, **kwargs):
        try:
            document_id = self.kwargs.get('document_id')
            queryset = self.queryset.filter(id=document_id)
            serializer = self.serializer_class(queryset, many=True, context={'request': request})
            return Response(serializer.data)
        except Exception as e:
            return Response({'status': False,'message': 'Error fetching documents', 'error': str(e)})
    